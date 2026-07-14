from __future__ import annotations

import argparse
import json
import subprocess
import sys
import unittest
from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document

from format_workflow import run_workflow


REPO_ROOT = Path(__file__).resolve().parents[1]


class FormatWorkflowTests(unittest.TestCase):
    def make_docx(self, path: Path) -> None:
        document = Document()
        document.add_paragraph("第一章 总则")
        document.add_paragraph("正文内容。")
        document.save(path)

    def make_structure(self, path: Path) -> None:
        path.write_text(
            json.dumps(
                {
                    "metadata": {"paragraph_count": 2, "table_count": 0},
                    "blocks": [],
                    "sections": [],
                    "preserve": [],
                },
                ensure_ascii=False,
                indent=2,
            ),
            encoding="utf-8",
        )

    def make_args(self, root: Path, **overrides) -> argparse.Namespace:
        values = {
            "format_package": None,
            "spec": root / "format_spec.json",
            "formatter": root / "formatter.py",
            "input": root / "raw.docx",
            "structure": root / "paper_structure.json",
            "output": root / "formatted.docx",
            "workflow_report_json": root / "workflow_report.json",
            "workflow_report_md": root / "workflow_report.md",
            "format_report_json": root / "format_report.json",
            "format_report_md": root / "format_report.md",
        }
        values.update(overrides)
        return argparse.Namespace(**values)

    def test_missing_spec_blocks_and_points_to_module_1(self) -> None:
        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            self.make_docx(root / "raw.docx")
            self.make_structure(root / "paper_structure.json")
            args = self.make_args(root, formatter=None)

            report = run_workflow(args)

            self.assertEqual(report["status"], "blocked")
            self.assertFalse((root / "formatted.docx").exists())
            self.assertIn("workflow_report.json", [p.name for p in root.iterdir()])
            missing = {item["artifact"]: item["return_to"] for item in report["missing"]}
            self.assertEqual(missing["format_spec.json"], "module_1")

    def test_missing_structure_blocks_and_points_to_module_3(self) -> None:
        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            self.make_docx(root / "raw.docx")
            (root / "format_spec.json").write_text('{"body": {}}', encoding="utf-8")
            args = self.make_args(root, formatter=None)

            report = run_workflow(args)

            self.assertEqual(report["status"], "blocked")
            missing = {item["artifact"]: item["return_to"] for item in report["missing"]}
            self.assertEqual(missing["paper_structure.json"], "module_3")

    def test_uses_formatter_when_available(self) -> None:
        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            self.make_docx(root / "raw.docx")
            self.make_structure(root / "paper_structure.json")
            (root / "format_spec.json").write_text('{"body": {}}', encoding="utf-8")
            formatter = root / "formatter.py"
            formatter.write_text(
                "from pathlib import Path\n"
                "import json, shutil, sys\n"
                "raw, out = Path(sys.argv[1]), Path(sys.argv[2])\n"
                "report = Path(sys.argv[sys.argv.index('--report') + 1])\n"
                "shutil.copyfile(raw, out)\n"
                "report.write_text(json.dumps({'status':'success','input':str(raw),'output':str(out),'applied':['mock']}), encoding='utf-8')\n",
                encoding="utf-8",
            )
            args = self.make_args(root)

            report = run_workflow(args)

            self.assertEqual(report["status"], "success")
            self.assertEqual(report["executor"], "formatter")
            self.assertTrue((root / "formatted.docx").is_file())
            self.assertTrue((root / "format_report.md").is_file())

    def test_falls_back_to_format_engine_without_formatter(self) -> None:
        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            self.make_docx(root / "raw.docx")
            self.make_structure(root / "paper_structure.json")
            (root / "format_spec.json").write_text(
                json.dumps(
                    {
                        "page": {"paper_size": "A4"},
                        "body": {
                            "font": {"east_asia": "宋体", "latin": "Times New Roman", "size_pt": 12},
                            "alignment": "justify",
                        },
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            args = self.make_args(root, formatter=None)

            report = run_workflow(args)

            self.assertEqual(report["status"], "success")
            self.assertEqual(report["executor"], "format_engine")
            self.assertTrue((root / "formatted.docx").is_file())
            self.assertTrue((root / "format_report.json").is_file())

    def test_formatter_failure_marks_workflow_failed(self) -> None:
        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            self.make_docx(root / "raw.docx")
            self.make_structure(root / "paper_structure.json")
            (root / "format_spec.json").write_text('{"body": {}}', encoding="utf-8")
            formatter = root / "formatter.py"
            formatter.write_text("import sys\nsys.exit(7)\n", encoding="utf-8")
            args = self.make_args(root)

            report = run_workflow(args)

            self.assertEqual(report["status"], "failed")
            self.assertTrue(report["errors"])
            self.assertFalse((root / "formatted.docx").exists())

    def test_creates_report_directories_before_engine_runs(self) -> None:
        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            self.make_docx(root / "raw.docx")
            self.make_structure(root / "paper_structure.json")
            (root / "format_spec.json").write_text('{"body": {}}', encoding="utf-8")
            args = self.make_args(
                root,
                formatter=None,
                output=root / "out" / "formatted.docx",
                workflow_report_json=root / "reports" / "workflow" / "workflow_report.json",
                workflow_report_md=root / "reports" / "workflow" / "workflow_report.md",
                format_report_json=root / "reports" / "format" / "format_report.json",
                format_report_md=root / "reports" / "format" / "format_report.md",
            )

            report = run_workflow(args)

            self.assertEqual(report["status"], "success")
            self.assertTrue((root / "out" / "formatted.docx").is_file())
            self.assertTrue((root / "reports" / "workflow" / "workflow_report.json").is_file())
            self.assertTrue((root / "reports" / "format" / "format_report.json").is_file())

    def test_cli_outputs_reports(self) -> None:
        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            self.make_docx(root / "raw.docx")
            self.make_structure(root / "paper_structure.json")
            (root / "format_spec.json").write_text('{"body": {}}', encoding="utf-8")

            subprocess.run(
                [
                    sys.executable,
                    str(REPO_ROOT / "format_workflow.py"),
                    "run",
                    "--input",
                    str(root / "raw.docx"),
                    "--spec",
                    str(root / "format_spec.json"),
                    "--structure",
                    str(root / "paper_structure.json"),
                    "--output",
                    str(root / "formatted.docx"),
                    "--workflow-report-json",
                    str(root / "workflow_report.json"),
                    "--workflow-report-md",
                    str(root / "workflow_report.md"),
                    "--format-report-json",
                    str(root / "format_report.json"),
                    "--format-report-md",
                    str(root / "format_report.md"),
                ],
                check=True,
                cwd=REPO_ROOT,
            )

            self.assertTrue((root / "workflow_report.json").is_file())
            self.assertTrue((root / "workflow_report.md").is_file())
            self.assertTrue((root / "format_report.json").is_file())
            self.assertTrue((root / "format_report.md").is_file())
            self.assertEqual(json.loads((root / "workflow_report.json").read_text(encoding="utf-8"))["status"], "success")


if __name__ == "__main__":
    unittest.main()
