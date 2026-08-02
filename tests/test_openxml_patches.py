from __future__ import annotations

import json
import unittest
import zipfile
from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree

from format_engine import format_document


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"


def xml_from_docx(path: Path, member: str):
    with zipfile.ZipFile(path, "r") as package:
        return etree.fromstring(package.read(member))


class OpenXmlPatchTests(unittest.TestCase):
    def make_docx(self, path: Path) -> None:
        document = Document()
        document.add_paragraph("测试文档标题")
        document.add_paragraph("图1-1 系统结构图")
        document.add_paragraph("表1-1 参数表")
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "参数"
        table.cell(0, 1).text = "值"
        table.cell(1, 0).text = "A"
        table.cell(1, 1).text = "1"
        document.add_paragraph("参考文献")
        document.add_paragraph("[1] 张三. 示例文献[J]. 2024.")
        document.add_paragraph("致谢")
        document.add_paragraph("E = mc^2 (1-1)")
        document.save(path)

    def base_spec(self) -> dict:
        return {
            "page": {"paper_size": "A4"},
            "body": {
                "font": {"east_asia": "宋体", "latin": "Times New Roman", "size_pt": 12},
                "alignment": "justify",
                "line_spacing": 1.5,
                "first_line_indent_chars": 2,
            },
        }

    def test_three_line_table_sets_table_and_header_borders(self) -> None:
        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            raw = root / "raw.docx"
            out = root / "out.docx"
            self.make_docx(raw)
            spec = self.base_spec()
            spec["tables"] = {"style": "three_line"}

            report = format_document(raw, out, spec)

            self.assertEqual(report["status"], "success")
            self.assertTrue(any(item.startswith("patch:three_line_table") for item in report["applied"]))
            document_xml = xml_from_docx(out, "word/document.xml")
            table_borders = document_xml.find(f".//{{{W_NS}}}tblBorders")
            self.assertIsNotNone(table_borders)
            self.assertEqual(table_borders.find(f"{{{W_NS}}}top").get(f"{{{W_NS}}}sz"), "12")
            self.assertEqual(table_borders.find(f"{{{W_NS}}}bottom").get(f"{{{W_NS}}}sz"), "12")
            self.assertEqual(table_borders.find(f"{{{W_NS}}}insideV").get(f"{{{W_NS}}}val"), "nil")
            cell_bottom = document_xml.find(f".//{{{W_NS}}}tcBorders/{{{W_NS}}}bottom")
            self.assertIsNotNone(cell_bottom)
            self.assertEqual(cell_bottom.get(f"{{{W_NS}}}sz"), "6")

    def test_header_footer_creates_parts_and_relationships(self) -> None:
        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            raw = root / "raw.docx"
            out = root / "out.docx"
            self.make_docx(raw)
            spec = self.base_spec()
            spec["headers_footers"] = {
                "header": {"text": "测试页眉", "font": "宋体", "size_pt": 12, "alignment": "center"},
                "footer": {"page_number": True, "font": "Times New Roman", "size_pt": 9, "alignment": "center"},
            }

            report = format_document(raw, out, spec)

            self.assertEqual(report["status"], "success")
            self.assertTrue(any(item.startswith("patch:header_footer") for item in report["applied"]))
            with zipfile.ZipFile(out, "r") as package:
                names = set(package.namelist())
                header_names = [name for name in names if name.startswith("word/header") and name.endswith(".xml")]
                footer_names = [name for name in names if name.startswith("word/footer") and name.endswith(".xml")]
                self.assertTrue(header_names)
                self.assertTrue(footer_names)
                header_xml = etree.fromstring(package.read(header_names[0]))
                footer_xml = etree.fromstring(package.read(footer_names[0]))
                self.assertIn("测试页眉", "".join(header_xml.itertext()))
                self.assertIsNotNone(footer_xml.find(f".//{{{W_NS}}}instrText"))

            rels = xml_from_docx(out, "word/_rels/document.xml.rels")
            rel_types = [rel.get("Type") for rel in rels.findall(f"{{{REL_NS}}}Relationship")]
            self.assertIn("http://schemas.openxmlformats.org/officeDocument/2006/relationships/header", rel_types)
            self.assertIn("http://schemas.openxmlformats.org/officeDocument/2006/relationships/footer", rel_types)
            document_xml = xml_from_docx(out, "word/document.xml")
            self.assertIsNotNone(document_xml.find(f".//{{{W_NS}}}headerReference"))
            self.assertIsNotNone(document_xml.find(f".//{{{W_NS}}}footerReference"))

    def test_captions_references_math_and_equation_patches(self) -> None:
        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            raw = root / "raw.docx"
            out = root / "out.docx"
            report_path = root / "report.json"
            self.make_docx(raw)
            spec = self.base_spec()
            spec.update(
                {
                    "tables": {
                        "caption_alignment": "center",
                        "font": {"east_asia": "宋体", "latin": "Times New Roman", "size_pt": 10.5},
                    },
                    "figures": {
                        "caption_alignment": "center",
                        "font": {"east_asia": "宋体", "latin": "Times New Roman", "size_pt": 10.5},
                    },
                    "references": {"alignment": "justify", "indent": "hanging"},
                    "equations": {"font": "Times New Roman", "numbering": "right_aligned"},
                }
            )

            report = format_document(raw, out, spec, report_path)

            self.assertEqual(report["status"], "success")
            applied = json.loads(report_path.read_text(encoding="utf-8"))["applied"]
            self.assertTrue(any(item.startswith("patch:captions") for item in applied))
            self.assertTrue(any(item.startswith("patch:references") for item in applied))
            self.assertTrue(any(item.startswith("patch:math_font") for item in applied))
            self.assertTrue(any(item.startswith("patch:equation_numbering") for item in applied))

            document = Document(str(out))
            self.assertEqual(document.paragraphs[1].alignment, WD_ALIGN_PARAGRAPH.CENTER)
            self.assertEqual(document.paragraphs[2].alignment, WD_ALIGN_PARAGRAPH.CENTER)
            self.assertAlmostEqual(document.paragraphs[4].paragraph_format.first_line_indent.cm, -0.74, places=2)
            self.assertEqual(document.paragraphs[-1].alignment, WD_ALIGN_PARAGRAPH.CENTER)

            settings_xml = xml_from_docx(out, "word/settings.xml")
            math_font = settings_xml.find(f".//{{{M_NS}}}mathFont")
            self.assertIsNotNone(math_font)
            self.assertEqual(math_font.get(f"{{{M_NS}}}val"), "Times New Roman")


if __name__ == "__main__":
    unittest.main()
