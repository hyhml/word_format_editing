from __future__ import annotations

import json
import subprocess
import sys
import unittest
from pathlib import Path
from tempfile import TemporaryDirectory

from format_parser import parse_format_requirements, write_outputs


class FormatParserTests(unittest.TestCase):
    def test_txt_generates_readable_and_machine_specs(self) -> None:
        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            source = root / "format.txt"
            source.write_text(
                "武汉科技大学本科毕业论文2024版。\n"
                "页面采用A4纸，页边距上2.5cm，下2.5cm，左2.5cm，右2.5cm。\n"
                "正文小四号宋体，Times New Roman，1.5倍行距，首行缩进2字符，两端对齐。\n"
                "一级标题三号黑体居中。\n"
                "表格采用三线表，图题在图下方，表题在表上方，图表标题五号宋体居中。\n",
                encoding="utf-8",
            )

            spec, spec_md, report = parse_format_requirements(
                [source],
                description="武汉科技大学 本科毕业论文 2024",
            )

            self.assertIn("# 武汉科技大学 本科毕业论文 2024", spec_md)
            self.assertEqual(spec["page"]["paper_size"], "A4")
            self.assertEqual(spec["body"]["font"]["east_asia"], "宋体")
            self.assertEqual(spec["body"]["font"]["size_pt"], 12)
            self.assertEqual(spec["body"]["line_spacing"], 1.5)
            self.assertEqual(spec["body"]["alignment"], "justify")
            self.assertEqual(spec["tables"]["style"], "three_line")
            self.assertEqual(spec["figures"]["caption_position"], "below")
            self.assertEqual(spec["tables"]["caption_position"], "above")
            self.assertFalse(report["conflicts"])

    def test_docx_extracts_paragraphs_and_tables(self) -> None:
        with TemporaryDirectory() as tmp:
            try:
                from docx import Document
            except ImportError:
                self.skipTest("python-docx not installed")

            root = Path(tmp)
            source = root / "format.docx"
            document = Document()
            document.add_paragraph("正文小四号宋体，1.5倍行距，首行缩进2字符，两端对齐。")
            table = document.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "表格"
            table.cell(0, 1).text = "三线表"
            document.save(source)

            spec, _spec_md, report = parse_format_requirements([source])

            self.assertEqual(report["sources"][0]["type"], "docx")
            self.assertEqual(report["sources"][0]["table_count"], 1)
            self.assertEqual(spec["tables"]["style"], "three_line")
            self.assertEqual(spec["body"]["font"]["east_asia"], "宋体")

    def test_pdf_without_extractable_text_reports_warning(self) -> None:
        with TemporaryDirectory() as tmp:
            try:
                from pypdf import PdfWriter
            except ImportError:
                self.skipTest("pypdf not installed")

            root = Path(tmp)
            source = root / "blank.pdf"
            writer = PdfWriter()
            writer.add_blank_page(width=200, height=200)
            with source.open("wb") as handle:
                writer.write(handle)

            spec, _spec_md, report = parse_format_requirements([source])

            self.assertEqual(spec["metadata"]["sources"][0]["type"], "pdf")
            self.assertTrue(report["warnings"])
            self.assertIn("PDF 文本为空", report["warnings"][-1]["warning"])

    def test_conflicting_fields_are_not_silently_overwritten(self) -> None:
        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            first = root / "a.txt"
            second = root / "b.txt"
            first.write_text("正文小四号宋体，1.5倍行距。", encoding="utf-8")
            second.write_text("正文五号宋体，1.5倍行距。", encoding="utf-8")

            spec, _spec_md, report = parse_format_requirements([first, second])

            conflict_fields = {item["field"] for item in spec["conflicts"]}
            self.assertIn("body.font.size_cn", conflict_fields)
            self.assertIn("body.font.size_pt", conflict_fields)
            self.assertEqual(report["conflicts"], spec["conflicts"])

    def test_unknowns_are_reported_for_missing_required_fields(self) -> None:
        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            source = root / "minimal.txt"
            source.write_text("仅说明：格式按照学校要求执行。", encoding="utf-8")

            spec, _spec_md, report = parse_format_requirements([source])

            fields = {item["field"] for item in spec["unknowns"]}
            self.assertIn("body.font.size_pt", fields)
            self.assertIn("page.paper_size", fields)
            self.assertEqual(report["unknowns"], spec["unknowns"])

    def test_cli_writes_expected_outputs(self) -> None:
        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            source = root / "format.txt"
            output_dir = root / "out"
            source.write_text("A4。正文小四号宋体，1.5倍行距，首行缩进2字符，两端对齐。", encoding="utf-8")

            subprocess.run(
                [
                    sys.executable,
                    str(Path(__file__).resolve().parents[1] / "format_parser.py"),
                    "parse",
                    "--output-dir",
                    str(output_dir),
                    "--description",
                    "通用公文格式示例",
                    str(source),
                ],
                check=True,
            )

            self.assertTrue((output_dir / "format_spec.md").is_file())
            self.assertTrue((output_dir / "format_spec.json").is_file())
            self.assertTrue((output_dir / "parse_report.json").is_file())
            data = json.loads((output_dir / "format_spec.json").read_text(encoding="utf-8"))
            self.assertEqual(data["body"]["font"]["size_pt"], 12)


if __name__ == "__main__":
    unittest.main()
