from __future__ import annotations

import json
import subprocess
import sys
import unittest
import zipfile
from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document
from docx.shared import Cm, Pt
from lxml import etree

from format_engine import format_document
from format_validator import validate_document


REPO_ROOT = Path(__file__).resolve().parents[1]
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


class FormatValidatorTests(unittest.TestCase):
    def make_raw_docx(self, path: Path, with_table: bool = True) -> None:
        document = Document()
        document.add_paragraph("测试文档标题")
        document.add_paragraph("第一章 总则")
        document.add_paragraph("这是正文段落。")
        if with_table:
            table = document.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "列一"
            table.cell(0, 1).text = "列二"
            table.cell(1, 0).text = "值一"
            table.cell(1, 1).text = "值二"
        document.save(path)

    def make_spec(self) -> dict:
        return {
            "page": {
                "size": "A4",
                "orientation": "portrait",
                "margins_cm": {"top": 2.5, "bottom": 2.5, "left": 3.0, "right": 3.0},
            },
            "default": {
                "font_name": "Times New Roman",
                "east_asia_font_name": "宋体",
                "font_size_pt": 12,
                "bold": False,
                "alignment": "justify",
                "line_spacing": 1.5,
                "first_line_indent_cm": 0.74,
                "space_before_pt": 0,
                "space_after_pt": 0,
            },
            "rules": [
                {
                    "name": "封面标题",
                    "match": {"paragraph_index": 0},
                    "format": {
                        "east_asia_font_name": "黑体",
                        "font_size_pt": 18,
                        "bold": True,
                        "alignment": "center",
                        "line_spacing": 1.5,
                        "first_line_indent_cm": 0,
                        "space_before_pt": 0,
                        "space_after_pt": 12,
                    },
                },
                {
                    "name": "一级标题",
                    "match": {"regex": "^第[一二三四五六七八九十百]+章"},
                    "format": {
                        "east_asia_font_name": "黑体",
                        "font_size_pt": 16,
                        "bold": True,
                        "alignment": "left",
                        "line_spacing": 1.5,
                        "first_line_indent_cm": 0,
                        "space_before_pt": 12,
                        "space_after_pt": 6,
                    },
                },
            ],
            "tables": {
                "font_name": "Times New Roman",
                "east_asia_font_name": "宋体",
                "font_size_pt": 10.5,
                "bold_header_row": True,
                "alignment": "center",
                "cell_vertical_alignment": "center",
            },
        }

    def make_structure(self, path: Path, with_table: bool = True) -> None:
        blocks = [
            {"id": "p0000", "type": "body", "text": "测试文档标题", "source": {"kind": "paragraph", "index": 0}},
            {"id": "p0001", "type": "heading", "level": 1, "text": "第一章 总则", "source": {"kind": "paragraph", "index": 1}},
            {"id": "p0002", "type": "body", "text": "这是正文段落。", "source": {"kind": "paragraph", "index": 2}},
        ]
        tables = []
        if with_table:
            tables.append({"id": "tbl0000", "type": "table", "source": {"kind": "table", "index": 0}})
        path.write_text(
            json.dumps(
                {
                    "metadata": {"paragraph_count": 3, "table_count": len(tables)},
                    "blocks": blocks + tables,
                    "sections": [],
                    "figures": [],
                    "tables": tables,
                    "table_captions": [],
                    "references": [],
                    "preserve": [],
                },
                ensure_ascii=False,
                indent=2,
            ),
            encoding="utf-8",
        )

    def make_advanced_docx(self, path: Path) -> None:
        document = Document()
        document.add_paragraph("测试文档标题")
        document.add_paragraph("图1-1 系统结构图")
        document.add_paragraph("表1-1 参数表")
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "参数"
        table.cell(0, 1).text = "值"
        table.cell(1, 0).text = "A"
        table.cell(1, 1).text = "1"
        document.add_paragraph("参考文献")
        document.add_paragraph("[1] 张三. 示例文献[J]. 2024.")
        document.add_paragraph("致谢")
        document.add_paragraph("E = mc^2 (1-1)")
        document.save(path)

    def make_advanced_spec(self) -> dict:
        return {
            "page": {"paper_size": "A4"},
            "body": {
                "font": {"east_asia": "宋体", "latin": "Times New Roman", "size_pt": 12},
                "alignment": "justify",
                "line_spacing": 1.5,
                "first_line_indent_chars": 2,
            },
            "tables": {
                "style": "three_line",
                "caption_alignment": "center",
                "font": {"east_asia": "宋体", "latin": "Times New Roman", "size_pt": 10.5},
            },
            "figures": {
                "caption_alignment": "center",
                "font": {"east_asia": "宋体", "latin": "Times New Roman", "size_pt": 10.5},
            },
            "references": {"alignment": "justify", "indent": "hanging"},
            "headers_footers": {
                "header": {"text": "测试页眉", "font": "宋体", "size_pt": 12, "alignment": "center"},
                "footer": {"page_number": True, "font": "Times New Roman", "size_pt": 9, "alignment": "center"},
            },
            "equations": {"font": "Times New Roman", "numbering": "right_aligned"},
        }

    def make_advanced_structure(self, path: Path) -> None:
        blocks = [
            {"id": "p0000", "type": "body", "text": "测试文档标题", "source": {"kind": "paragraph", "index": 0}},
            {"id": "p0001", "type": "figure_caption", "text": "图1-1 系统结构图", "source": {"kind": "paragraph", "index": 1}},
            {"id": "p0002", "type": "table_caption", "text": "表1-1 参数表", "source": {"kind": "paragraph", "index": 2}},
            {"id": "p0004", "type": "reference_item", "text": "[1] 张三. 示例文献[J]. 2024.", "source": {"kind": "paragraph", "index": 4}},
            {"id": "tbl0000", "type": "table", "source": {"kind": "table", "index": 0}},
        ]
        path.write_text(
            json.dumps(
                {
                    "metadata": {"paragraph_count": 7, "table_count": 1},
                    "blocks": blocks,
                    "sections": [],
                    "figures": [blocks[1]],
                    "tables": [blocks[4]],
                    "table_captions": [blocks[2]],
                    "references": [blocks[3]],
                    "preserve": [],
                },
                ensure_ascii=False,
                indent=2,
            ),
            encoding="utf-8",
        )

    def make_advanced_formatted_case(self, root: Path) -> tuple[Path, Path, Path]:
        raw = root / "raw.docx"
        formatted = root / "formatted.docx"
        spec_path = root / "format_spec.json"
        structure_path = root / "paper_structure.json"
        self.make_advanced_docx(raw)
        spec = self.make_advanced_spec()
        spec_path.write_text(json.dumps(spec, ensure_ascii=False, indent=2), encoding="utf-8")
        self.make_advanced_structure(structure_path)
        format_document(raw, formatted, spec)
        return formatted, spec_path, structure_path

    def make_formatted_case(self, root: Path, with_table: bool = True) -> tuple[Path, Path, Path]:
        raw = root / "raw.docx"
        formatted = root / "formatted.docx"
        spec_path = root / "format_spec.json"
        structure_path = root / "paper_structure.json"
        self.make_raw_docx(raw, with_table=with_table)
        spec = self.make_spec()
        spec_path.write_text(json.dumps(spec, ensure_ascii=False, indent=2), encoding="utf-8")
        self.make_structure(structure_path, with_table=with_table)
        format_document(raw, formatted, spec)
        return formatted, spec_path, structure_path

    def test_compliant_document_passes(self) -> None:
        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            formatted, spec_path, structure_path = self.make_formatted_case(root)

            report = validate_document(formatted, spec_path, structure_path)

            self.assertEqual(report["status"], "pass")
            self.assertGreater(report["summary"]["pass"], 0)
            self.assertEqual(report["summary"]["fail"], 0)

    def test_wrong_page_margin_fails(self) -> None:
        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            formatted, spec_path, structure_path = self.make_formatted_case(root)
            document = Document(str(formatted))
            document.sections[0].top_margin = Cm(3.5)
            document.save(str(formatted))

            report = validate_document(formatted, spec_path, structure_path)

            self.assertEqual(report["status"], "fail")
            self.assertTrue(any(item["id"] == "page.margin.top" and item["result"] == "fail" for item in report["checks"]))

    def test_wrong_body_font_size_fails(self) -> None:
        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            formatted, spec_path, structure_path = self.make_formatted_case(root)
            document = Document(str(formatted))
            document.paragraphs[2].runs[0].font.size = Pt(9)
            document.save(str(formatted))

            report = validate_document(formatted, spec_path, structure_path)

            self.assertEqual(report["status"], "fail")
            self.assertTrue(any(item["id"] == "legacy.default.font.size_pt" and item["result"] == "fail" for item in report["checks"]))

    def test_unknown_spec_field_warns(self) -> None:
        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            formatted, spec_path, structure_path = self.make_formatted_case(root)
            spec = json.loads(spec_path.read_text(encoding="utf-8"))
            spec["default"]["font_size_pt"] = "unknown"
            spec_path.write_text(json.dumps(spec, ensure_ascii=False, indent=2), encoding="utf-8")

            report = validate_document(formatted, spec_path, structure_path)

            self.assertEqual(report["status"], "warn")
            self.assertTrue(any(item["id"] == "legacy.default.font.size_pt" and item["result"] == "warn" for item in report["checks"]))

    def test_missing_tables_warns(self) -> None:
        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            formatted, spec_path, structure_path = self.make_formatted_case(root, with_table=False)

            report = validate_document(formatted, spec_path, structure_path)

            self.assertEqual(report["status"], "warn")
            self.assertTrue(any(item["id"] == "tables.exists" and item["result"] == "warn" for item in report["checks"]))

    def test_cli_writes_validation_reports(self) -> None:
        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            formatted, spec_path, structure_path = self.make_formatted_case(root)
            report_json = root / "validation_report.json"
            report_md = root / "validation_report.md"

            subprocess.run(
                [
                    sys.executable,
                    str(REPO_ROOT / "format_validator.py"),
                    "validate",
                    "--input",
                    str(formatted),
                    "--spec",
                    str(spec_path),
                    "--structure",
                    str(structure_path),
                    "--report-json",
                    str(report_json),
                    "--report-md",
                    str(report_md),
                ],
                check=True,
                cwd=REPO_ROOT,
            )

            self.assertTrue(report_json.is_file())
            self.assertTrue(report_md.is_file())
            self.assertEqual(json.loads(report_json.read_text(encoding="utf-8"))["status"], "pass")

    def test_advanced_openxml_rules_validate(self) -> None:
        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            formatted, spec_path, structure_path = self.make_advanced_formatted_case(root)

            report = validate_document(formatted, spec_path, structure_path)

            self.assertEqual(report["status"], "warn")
            self.assertEqual(report["summary"]["fail"], 0)
            check_ids = {item["id"] for item in report["checks"] if item["result"] == "pass"}
            self.assertIn("table.openxml.top.val", check_ids)
            self.assertIn("headers_footers.header.text", check_ids)
            self.assertIn("headers_footers.footer.page_number", check_ids)
            self.assertIn("references.first_line_indent", check_ids)
            self.assertIn("equations.math_font.settings", check_ids)
            self.assertTrue(any(item["id"] == "equations.numbering.right_aligned" and item["result"] == "warn" for item in report["checks"]))

    def test_broken_three_line_table_fails(self) -> None:
        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            formatted, spec_path, structure_path = self.make_advanced_formatted_case(root)
            broken = root / "broken.docx"
            with zipfile.ZipFile(formatted, "r") as source, zipfile.ZipFile(broken, "w", zipfile.ZIP_DEFLATED) as target:
                for item in source.infolist():
                    data = source.read(item.filename)
                    if item.filename == "word/document.xml":
                        xml = etree.fromstring(data)
                        top = xml.find(f".//{{{W_NS}}}tblBorders/{{{W_NS}}}top")
                        top.set(f"{{{W_NS}}}sz", "4")
                        data = etree.tostring(xml, encoding="UTF-8", xml_declaration=True)
                    target.writestr(item, data)

            report = validate_document(broken, spec_path, structure_path)

            self.assertEqual(report["status"], "fail")
            self.assertTrue(any(item["id"] == "table.openxml.top.sz" and item["result"] == "fail" for item in report["checks"]))

    def test_broken_header_text_fails(self) -> None:
        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            formatted, spec_path, structure_path = self.make_advanced_formatted_case(root)
            spec = json.loads(spec_path.read_text(encoding="utf-8"))
            spec["headers_footers"]["header"]["text"] = "不存在的页眉"
            spec_path.write_text(json.dumps(spec, ensure_ascii=False, indent=2), encoding="utf-8")

            report = validate_document(formatted, spec_path, structure_path)

            self.assertEqual(report["status"], "fail")
            self.assertTrue(any(item["id"] == "headers_footers.header.text" and item["result"] == "fail" for item in report["checks"]))


if __name__ == "__main__":
    unittest.main()
