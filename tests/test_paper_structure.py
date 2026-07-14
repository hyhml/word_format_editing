from __future__ import annotations

import json
import subprocess
import sys
import unittest
from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document

from paper_structure import analyze_docx


REPO_ROOT = Path(__file__).resolve().parents[1]


class PaperStructureTests(unittest.TestCase):
    def make_sample_docx(self, path: Path) -> None:
        document = Document()
        document.add_paragraph("摘要")
        document.add_paragraph("本文研究一个示例问题。")
        document.add_paragraph("关键词：格式；结构；识别")
        document.add_heading("第一章 绪论", level=1)
        document.add_paragraph("1.1 研究背景")
        document.add_paragraph("这是一个正文段落。")
        document.add_paragraph("图1-1 系统结构图")
        document.add_paragraph("表1-1 参数表")
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "参数"
        table.cell(0, 1).text = "值"
        table.cell(1, 0).text = "A"
        table.cell(1, 1).text = "1"
        document.add_paragraph("参考文献")
        document.add_paragraph("[1] 张三. 示例文献[J]. 2024.")
        document.add_paragraph("[2] 李四. 示例文献[M]. 2023.")
        document.add_paragraph("致谢")
        document.save(path)

    def test_analyze_docx_identifies_core_structure(self) -> None:
        with TemporaryDirectory() as tmp:
            source = Path(tmp) / "paper.docx"
            self.make_sample_docx(source)

            structure, markdown, report = analyze_docx(source)

            self.assertIn("# 论文结构识别结果", markdown)
            self.assertEqual(report["status"], "success")
            self.assertEqual(report["counts"]["tables"], 1)
            self.assertGreaterEqual(report["counts"]["headings"], 2)
            self.assertEqual(report["counts"]["figures"], 1)
            self.assertEqual(report["counts"]["table_captions"], 1)
            self.assertEqual(report["counts"]["references"], 2)

            section_titles = [section["title"] for section in structure["sections"]]
            self.assertIn("第一章 绪论", section_titles)
            self.assertIn("1.1 研究背景", section_titles)
            self.assertEqual(structure["figures"][0]["number"], "1-1")
            self.assertEqual(structure["table_captions"][0]["number"], "1-1")
            self.assertIn("source", structure["blocks"][0])

    def test_unclassified_body_is_preserved(self) -> None:
        with TemporaryDirectory() as tmp:
            source = Path(tmp) / "paper.docx"
            document = Document()
            document.add_paragraph("普通正文，没有标题模式。")
            document.save(source)

            structure, _markdown, report = analyze_docx(source)

            self.assertEqual(report["counts"]["preserve"], 1)
            self.assertEqual(structure["preserve"][0]["reason"], "unclassified_paragraph")
            self.assertEqual(structure["preserve"][0]["source"]["index"], 0)

    def test_style_heading_has_higher_confidence(self) -> None:
        with TemporaryDirectory() as tmp:
            source = Path(tmp) / "paper.docx"
            document = Document()
            document.add_heading("无编号标题", level=1)
            document.save(source)

            structure, _markdown, _report = analyze_docx(source)

            self.assertEqual(structure["sections"][0]["level"], 1)
            self.assertEqual(structure["sections"][0]["method"], "style")
            self.assertEqual(structure["sections"][0]["confidence"], 0.95)

    def test_cli_writes_structure_outputs(self) -> None:
        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            source = root / "paper.docx"
            output_md = root / "paper_structure.md"
            output_json = root / "paper_structure.json"
            report = root / "structure_report.json"
            self.make_sample_docx(source)

            subprocess.run(
                [
                    sys.executable,
                    str(REPO_ROOT / "paper_structure.py"),
                    "analyze",
                    "--input",
                    str(source),
                    "--output-md",
                    str(output_md),
                    "--output-json",
                    str(output_json),
                    "--report",
                    str(report),
                ],
                check=True,
                cwd=REPO_ROOT,
            )

            self.assertTrue(output_md.is_file())
            self.assertTrue(output_json.is_file())
            self.assertTrue(report.is_file())
            data = json.loads(output_json.read_text(encoding="utf-8"))
            self.assertIn("blocks", data)
            self.assertIn("sections", data)
            self.assertTrue(data["sections"])


if __name__ == "__main__":
    unittest.main()
