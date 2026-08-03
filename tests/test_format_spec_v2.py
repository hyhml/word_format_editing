from __future__ import annotations

import json
import re
import subprocess
import sys
import unittest
from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document

from format_compiler import compile_sources, empty_spec
from format_normalization import normalize_font_size, normalize_length_cm, normalize_line_spacing
from format_ontology import PROPERTY_NAMES, TARGETS
from format_spec_validator_v2 import load_and_validate, validate_spec
from requirement_source import extract_requirement_sources
from requirement_patterns import PATTERNS
from selector_patterns import SELECTOR_PATTERNS


REPO_ROOT = Path(__file__).resolve().parents[1]


class FormatSpecV2Tests(unittest.TestCase):
    def minimal_spec(self) -> dict:
        with TemporaryDirectory() as tmp:
            source = Path(tmp) / "requirements.txt"
            source.write_text("没有具体要求。", encoding="utf-8")
            sources = extract_requirement_sources([source])
            return empty_spec(sources, "测试规范", "")

    def test_ontology_has_core_targets_and_properties(self) -> None:
        self.assertIn("body.paragraph", TARGETS)
        self.assertIn("equation.number", TARGETS)
        self.assertIn("page.margin_top_cm", PROPERTY_NAMES)
        self.assertIn("paragraph.keep_with_next", PROPERTY_NAMES)

    def test_minimal_spec_is_schema_valid(self) -> None:
        report = validate_spec(self.minimal_spec())
        self.assertEqual(report["status"], "success")
        self.assertFalse(report["errors"])

    def test_invalid_value_and_regex_return_repair_feedback(self) -> None:
        spec = self.minimal_spec()
        spec["targets"]["body.paragraph"] = {
            "properties": {
                "font.size_pt": {
                    "action": "set",
                    "value": -12,
                    "unit": "cm",
                    "evidence_ids": ["source_01_line_0001"],
                }
            }
        }
        spec["selectors"]["body.paragraph"] = {"fallback_regex": ["("]}

        report = validate_spec(spec)

        self.assertEqual(report["status"], "failed")
        error_types = {item["error_type"] for item in report["errors"]}
        self.assertIn("out_of_range", error_types)
        self.assertIn("invalid_unit", error_types)
        self.assertIn("invalid_regex", error_types)
        self.assertEqual(report["repair_request"]["errors"], report["errors"])

    def test_preserve_unresolved_is_valid_but_warns(self) -> None:
        spec = self.minimal_spec()
        spec["targets"]["body.paragraph"] = {
            "properties": {"font.size_pt": {"action": "preserve", "reason": "unresolved_conflict"}}
        }

        report = validate_spec(spec)

        self.assertEqual(report["status"], "warn")
        self.assertFalse(report["errors"])

    def test_duplicate_json_key_is_rejected_before_validation(self) -> None:
        with TemporaryDirectory() as tmp:
            path = Path(tmp) / "format_spec.json"
            path.write_text('{"schema_version":"2.0.0","schema_version":"2.0.0"}', encoding="utf-8")
            with self.assertRaisesRegex(ValueError, "重复JSON字段"):
                load_and_validate(path)

    def test_normalizes_common_units(self) -> None:
        self.assertEqual(normalize_font_size("小四号"), 12)
        self.assertEqual(normalize_font_size("10.5pt"), 10.5)
        self.assertEqual(normalize_length_cm("25毫米"), 2.5)
        self.assertAlmostEqual(normalize_length_cm("1英寸"), 2.54)
        self.assertEqual(normalize_line_spacing("固定值20磅"), ("fixed", 20.0))

    def test_selector_patterns_match_positive_not_negative_examples(self) -> None:
        for target, definition in SELECTOR_PATTERNS.items():
            patterns = [re.compile(pattern, re.IGNORECASE) for pattern in definition["patterns"]]
            for example in definition["positive"]:
                self.assertTrue(any(pattern.search(example) for pattern in patterns), f"{target}: {example}")
            for example in definition["negative"]:
                self.assertFalse(any(pattern.search(example) for pattern in patterns), f"{target}: {example}")

    def test_all_requirement_patterns_compile(self) -> None:
        for pattern in PATTERNS:
            self.assertIsNotNone(pattern.compile())


class RequirementSourceTests(unittest.TestCase):
    def test_docx_preserves_paragraph_style_and_table_location(self) -> None:
        with TemporaryDirectory() as tmp:
            path = Path(tmp) / "requirements.docx"
            document = Document()
            paragraph = document.add_paragraph("正文小四号宋体")
            paragraph.style = document.styles["Normal"]
            paragraph.add_run("加粗").bold = True
            table = document.add_table(rows=1, cols=1)
            table.cell(0, 0).text = "三线表"
            document.save(path)

            source = extract_requirement_sources([path])[0]

            paragraph_block = next(block for block in source.blocks if block.type == "paragraph")
            table_block = next(block for block in source.blocks if block.type == "table_cell")
            section_block = next(block for block in source.blocks if block.type == "document_settings")
            self.assertEqual(paragraph_block.paragraph_index, 0)
            self.assertEqual(paragraph_block.style["style_name"], "Normal")
            self.assertEqual(table_block.table_index, 0)
            self.assertEqual((table_block.row, table_block.column), (0, 0))
            self.assertIn("page_width_cm", section_block.style)


class FormatCompilerTests(unittest.TestCase):
    def test_compiles_common_requirements_to_canonical_spec(self) -> None:
        with TemporaryDirectory() as tmp:
            source = Path(tmp) / "requirements.txt"
            source.write_text(
                "页面采用A4纸，上边距25毫米，下边距2.5厘米，左边距2.5cm，右边距2.5cm。\n"
                "正文小四号宋体，Times New Roman，1.5倍行距，首行缩进2字符，两端对齐。\n"
                "一级标题三号黑体居中加粗。\n"
                "表格采用三线表，表题在表上方，图题在图下方。\n",
                encoding="utf-8",
            )

            spec, report, _sources = compile_sources([source], "测试论文格式", "")

            self.assertNotEqual(report["status"], "blocked")
            self.assertEqual(spec["schema_version"], "2.0.0")
            self.assertEqual(spec["document"]["properties"]["page.margin_top_cm"]["value"], 2.5)
            body = spec["targets"]["body.paragraph"]["properties"]
            self.assertEqual(body["font.size_pt"]["value"], 12)
            self.assertEqual(body["paragraph.line_spacing.value"]["value"], 1.5)
            self.assertEqual(spec["targets"]["table"]["properties"]["table.style"]["value"], "three_line")
            self.assertFalse(report["validation"]["errors"])

    def test_conflict_becomes_preserve_in_unique_spec(self) -> None:
        with TemporaryDirectory() as tmp:
            first = Path(tmp) / "a.txt"
            second = Path(tmp) / "b.txt"
            first.write_text("正文小四号宋体。", encoding="utf-8")
            second.write_text("正文五号宋体。", encoding="utf-8")

            spec, report, _sources = compile_sources([first, second])

            size = spec["targets"]["body.paragraph"]["properties"]["font.size_pt"]
            self.assertEqual(size["action"], "preserve")
            self.assertEqual(size["reason"], "unresolved_conflict")
            self.assertTrue(report["conflicts"])
            self.assertEqual(report["validation"]["status"], "warn")

    def test_ai_candidate_can_fill_semantic_rule_with_evidence(self) -> None:
        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            source = root / "requirements.txt"
            source.write_text("一级标题同摘要标题，但不编号。", encoding="utf-8")
            ai = root / "ai.json"
            ai.write_text(
                json.dumps(
                    {
                        "candidates": [
                            {
                                "target": "heading.level_1",
                                "property": "numbering.enabled",
                                "value": False,
                                "evidence_ids": ["source_01_line_0001"],
                                "confidence": 0.9,
                            }
                        ]
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            spec, report, _sources = compile_sources([source], ai_candidates_path=ai)

            action = spec["targets"]["heading.level_1"]["properties"]["numbering.enabled"]
            self.assertEqual(action["action"], "set")
            self.assertEqual(action["method"], "ai")
            self.assertEqual(report["ai_candidate_count"], 1)

    def test_ai_candidate_with_unknown_evidence_is_rejected(self) -> None:
        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            source = root / "requirements.txt"
            source.write_text("正文格式另行规定。", encoding="utf-8")
            ai = root / "ai.json"
            ai.write_text(
                json.dumps(
                    {
                        "candidates": [
                            {
                                "target": "body.paragraph",
                                "property": "font.size_pt",
                                "value": 12,
                                "unit": "pt",
                                "evidence_ids": ["invented_evidence"],
                            }
                        ]
                    }
                ),
                encoding="utf-8",
            )

            spec, report, _sources = compile_sources([source], ai_candidates_path=ai)

            self.assertNotIn("body.paragraph", spec["targets"])
            self.assertEqual(report["ai_candidate_count"], 0)
            self.assertTrue(report["ai_candidate_errors"])

    def test_ai_classification_closes_non_rule_coverage_without_inventing_format(self) -> None:
        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            source = root / "requirements.txt"
            source.write_text("本文件用于说明论文提交注意事项。", encoding="utf-8")
            ai = root / "ai.json"
            ai.write_text(
                json.dumps(
                    {
                        "candidates": [],
                        "block_classifications": [
                            {
                                "evidence_id": "source_01_line_0001",
                                "classification": "explanation",
                                "notes": "不包含可执行格式要求",
                            }
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            _spec, report, _sources = compile_sources([source], ai_candidates_path=ai)

            self.assertFalse(report["coverage"]["unmapped_blocks"])
            self.assertFalse(report["coverage"]["unresolved_blocks"])
            self.assertEqual(report["coverage"]["mapped_ratio"], 1.0)

    def test_requirement_classification_without_rule_remains_unresolved(self) -> None:
        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            source = root / "requirements.txt"
            source.write_text("其他格式按照附件规定。", encoding="utf-8")
            ai = root / "ai.json"
            ai.write_text(
                json.dumps(
                    {
                        "candidates": [],
                        "block_classifications": [
                            {"evidence_id": "source_01_line_0001", "classification": "requirement"}
                        ],
                    }
                ),
                encoding="utf-8",
            )

            _spec, report, _sources = compile_sources([source], ai_candidates_path=ai)

            self.assertEqual(report["coverage"]["unresolved_blocks"], ["source_01_line_0001"])
            self.assertEqual(report["status"], "warn")

    def test_cli_writes_only_three_official_outputs(self) -> None:
        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            source = root / "requirements.txt"
            output = root / "out"
            source.write_text("正文小四号宋体。", encoding="utf-8")

            subprocess.run(
                [
                    sys.executable,
                    str(REPO_ROOT / "format_compiler.py"),
                    "compile",
                    "--source",
                    str(source),
                    "--output-dir",
                    str(output),
                ],
                cwd=REPO_ROOT,
                check=True,
            )

            self.assertEqual(
                {path.name for path in output.iterdir()},
                {"format_spec.json", "recognition_report.json", "recognition_report.md"},
            )

    def test_ai_request_contains_ontology_blocks_and_classification_contract(self) -> None:
        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            source = root / "requirements.txt"
            output = root / "ai_request.json"
            source.write_text("正文小四号宋体。", encoding="utf-8")

            subprocess.run(
                [
                    sys.executable,
                    str(REPO_ROOT / "format_compiler.py"),
                    "ai-request",
                    "--source",
                    str(source),
                    "--output",
                    str(output),
                ],
                cwd=REPO_ROOT,
                check=True,
            )
            request = json.loads(output.read_text(encoding="utf-8"))
            self.assertTrue(request["ontology"]["targets"])
            self.assertEqual(request["blocks"][0]["id"], "source_01_line_0001")
            self.assertIn("block_classifications", request["output_contract"])


if __name__ == "__main__":
    unittest.main()
