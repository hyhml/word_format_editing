from __future__ import annotations

import json
import subprocess
import sys
import unittest
from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document

from format_engine import format_document, load_spec


REPO_ROOT = Path(__file__).resolve().parents[1]


class FormatEngineTests(unittest.TestCase):
    def make_docx(self, path: Path) -> None:
        document = Document()
        document.add_paragraph("测试文档标题")
        document.add_paragraph("第一章 总则")
        document.add_paragraph("这是正文段落。")
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "列一"
        table.cell(0, 1).text = "列二"
        table.cell(1, 0).text = "值一"
        table.cell(1, 1).text = "值二"
        document.save(path)

    def test_legacy_spec_formats_docx(self) -> None:
        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            raw = root / "raw.docx"
            out = root / "out.docx"
            report_path = root / "report.json"
            self.make_docx(raw)

            spec = load_spec(REPO_ROOT / "examples" / "format_spec.example.json")
            report = format_document(raw, out, spec, report_path, REPO_ROOT / "examples" / "format_spec.example.json")

            self.assertEqual(report["status"], "success")
            self.assertTrue(out.is_file())
            saved_report = json.loads(report_path.read_text(encoding="utf-8"))
            self.assertIn("page", saved_report["applied"])
            self.assertTrue(any(item.startswith("body:") for item in saved_report["applied"]))
            self.assertTrue(any(item.startswith("tables:") for item in saved_report["applied"]))

    def test_module1_spec_formats_body_heading_and_tables(self) -> None:
        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            raw = root / "raw.docx"
            out = root / "out.docx"
            self.make_docx(raw)
            spec = {
                "page": {
                    "paper_size": "A4",
                    "orientation": "portrait",
                    "margins_cm": {"top": 2.5, "bottom": 2.5, "left": 2.5, "right": 2.5},
                },
                "body": {
                    "font": {"east_asia": "宋体", "latin": "Times New Roman", "size_pt": 12},
                    "alignment": "justify",
                    "line_spacing": 1.5,
                    "first_line_indent_chars": 2,
                },
                "headings": [
                    {
                        "level": 1,
                        "font": {"east_asia": "黑体", "latin": "Times New Roman", "size_pt": 16},
                        "alignment": "left",
                        "bold": True,
                    }
                ],
                "tables": {
                    "font": {"east_asia": "宋体", "latin": "Times New Roman", "size_pt": 10.5},
                    "caption_alignment": "center",
                    "bold_header_row": True,
                },
            }

            report = format_document(raw, out, spec)

            self.assertEqual(report["status"], "success")
            self.assertTrue(out.is_file())
            self.assertIn("page", report["applied"])
            self.assertTrue(any(item.startswith("headings:") for item in report["applied"]))
            self.assertTrue(any(item.startswith("tables:") for item in report["applied"]))

    def test_unknown_fields_are_skipped_in_report(self) -> None:
        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            raw = root / "raw.docx"
            out = root / "out.docx"
            self.make_docx(raw)
            spec = {
                "page": {"paper_size": "unknown", "margins_cm": {"top": "unknown"}},
                "body": {"font": {"east_asia": "unknown", "size_pt": "unknown"}},
                "tables": {},
                "openxml_patches": ["header_footer", "not_a_patch"],
            }

            report = format_document(raw, out, spec)

            self.assertEqual(report["status"], "success")
            self.assertTrue(out.is_file())
            self.assertIn("page: no known values", report["skipped"])
            self.assertIn("header_footer: not implemented in module 2", report["skipped_patches"])
            self.assertIn("not_a_patch: unknown patch", report["skipped_patches"])

    def test_failed_formatting_does_not_create_output(self) -> None:
        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            missing = root / "missing.docx"
            out = root / "out.docx"
            report_path = root / "report.json"

            report = format_document(missing, out, {"body": {}}, report_path)

            self.assertEqual(report["status"], "failed")
            self.assertFalse(out.exists())
            saved_report = json.loads(report_path.read_text(encoding="utf-8"))
            self.assertEqual(saved_report["status"], "failed")
            self.assertTrue(saved_report["errors"])

    def test_generated_formatter_wrapper_calls_engine(self) -> None:
        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            raw = root / "raw.docx"
            out = root / "out.docx"
            report = root / "report.json"
            wrapper = root / "formatter.py"
            self.make_docx(raw)

            subprocess.run(
                [
                    sys.executable,
                    str(REPO_ROOT / "generate_formatter.py"),
                    "--spec",
                    str(REPO_ROOT / "examples" / "format_spec.example.json"),
                    "--output",
                    str(wrapper),
                ],
                check=True,
                cwd=REPO_ROOT,
            )
            subprocess.run(
                [sys.executable, str(wrapper), str(raw), str(out), "--report", str(report)],
                check=True,
                cwd=REPO_ROOT,
            )

            self.assertTrue(out.is_file())
            self.assertEqual(json.loads(report.read_text(encoding="utf-8"))["status"], "success")


if __name__ == "__main__":
    unittest.main()
