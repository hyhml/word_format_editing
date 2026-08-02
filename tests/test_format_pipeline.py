from __future__ import annotations

import argparse
import json
import subprocess
import sys
import unittest
from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document

from format_pipeline import run_pipeline


REPO_ROOT = Path(__file__).resolve().parents[1]


class FormatPipelineTests(unittest.TestCase):
    def make_docx(self, path: Path) -> None:
        document = Document()
        document.add_paragraph("测试文档标题")
        document.add_paragraph("第一章 总则")
        document.add_paragraph("这是正文段落。")
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "列一"
        table.cell(0, 1).text = "列二"
        table.cell(1, 0).text = "值一"
        table.cell(1, 1).text = "值二"
        document.save(path)

    def make_format_source(self, path: Path) -> None:
        path.write_text(
            "页面采用A4纸，页边距上2.5cm，下2.5cm，左2.5cm，右2.5cm。\n"
            "正文小四号宋体，Times New Roman，1.5倍行距，首行缩进2字符，两端对齐。\n"
            "一级标题三号黑体左对齐加粗。\n"
            "表格采用三线表，图题在图下方，表题在表上方，图表标题五号宋体居中。\n",
            encoding="utf-8",
        )

    def make_args(self, root: Path, output_dir: Path | None = None) -> argparse.Namespace:
        return argparse.Namespace(
            command="run",
            input=root / "raw.docx",
            format_source=[root / "format.txt"],
            output_dir=output_dir or root / "out",
            formats_dir=root / "formats",
            description="通用论文格式 2026",
            name="通用论文格式 2026",
            package_id="demo_format",
            metadata_threshold=0.6,
        )

    def test_pipeline_creates_package_and_outputs_reports(self) -> None:
        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            self.make_docx(root / "raw.docx")
            self.make_format_source(root / "format.txt")

            report = run_pipeline(self.make_args(root))

            self.assertIn(report["status"], {"success", "warn"})
            self.assertEqual(report["package_action"], "created")
            self.assertEqual(report["workflow_status"], "success")
            self.assertTrue((root / "formats" / "demo_format" / "manifest.json").is_file())
            self.assertTrue((root / "formats" / "demo_format" / "formatter.py").is_file())
            self.assertTrue((root / "out" / "formatted.docx").is_file())
            self.assertTrue((root / "out" / "pipeline_report.json").is_file())
            self.assertTrue((root / "out" / "validation_report.json").is_file())

    def test_pipeline_reuses_existing_package_on_second_run(self) -> None:
        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            self.make_docx(root / "raw.docx")
            self.make_format_source(root / "format.txt")

            first = run_pipeline(self.make_args(root, root / "out_first"))
            second = run_pipeline(self.make_args(root, root / "out_second"))

            self.assertEqual(first["package_action"], "created")
            self.assertEqual(second["package_action"], "reused")
            self.assertEqual(second["registry"]["match_type"], "exact_hash")
            self.assertTrue((root / "out_second" / "formatted.docx").is_file())

    def test_cli_runs_full_pipeline(self) -> None:
        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            self.make_docx(root / "raw.docx")
            self.make_format_source(root / "format.txt")

            subprocess.run(
                [
                    sys.executable,
                    str(REPO_ROOT / "format_pipeline.py"),
                    "run",
                    "--input",
                    str(root / "raw.docx"),
                    "--format-source",
                    str(root / "format.txt"),
                    "--output-dir",
                    str(root / "out"),
                    "--formats-dir",
                    str(root / "formats"),
                    "--description",
                    "通用论文格式 2026",
                    "--name",
                    "通用论文格式 2026",
                    "--package-id",
                    "demo_format",
                ],
                check=True,
                cwd=REPO_ROOT,
            )

            report = json.loads((root / "out" / "pipeline_report.json").read_text(encoding="utf-8"))
            self.assertIn(report["status"], {"success", "warn"})
            self.assertEqual(report["workflow_status"], "success")
            self.assertTrue((root / "out" / "formatted.docx").is_file())


if __name__ == "__main__":
    unittest.main()
