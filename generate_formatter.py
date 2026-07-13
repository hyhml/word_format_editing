#!/usr/bin/env python3
"""Generate a standalone Word formatter from a JSON format specification."""

from __future__ import annotations

import argparse
import json
from pathlib import Path
from textwrap import dedent


REQUIRED_TOP_LEVEL_KEYS = {"default"}


def load_spec(path: Path) -> dict:
    with path.open("r", encoding="utf-8") as handle:
        spec = json.load(handle)

    missing = REQUIRED_TOP_LEVEL_KEYS - set(spec)
    if missing:
        raise ValueError(f"规格文件缺少必要字段: {', '.join(sorted(missing))}")

    if not isinstance(spec.get("rules", []), list):
        raise ValueError("字段 rules 必须是数组")

    for index, rule in enumerate(spec.get("rules", [])):
        if "match" not in rule or "format" not in rule:
            raise ValueError(f"rules[{index}] 必须包含 match 和 format")

    return spec


def build_formatter_source(spec: dict) -> str:
    spec_literal = json.dumps(spec, ensure_ascii=False, indent=2)
    template = r'''
#!/usr/bin/env python3
"""Standalone Word formatter generated from a JSON specification."""

from __future__ import annotations

import argparse
import copy
import json
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


SPEC = json.loads(__SPEC_JSON__)


ALIGNMENTS = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


TABLE_VERTICAL_ALIGNMENTS = {
    "top": WD_ALIGN_VERTICAL.TOP,
    "center": WD_ALIGN_VERTICAL.CENTER,
    "bottom": WD_ALIGN_VERTICAL.BOTTOM,
}


PAGE_SIZES_CM = {
    "A4": (21.0, 29.7),
    "LETTER": (21.59, 27.94),
}


def merged_format(base: dict, override: dict | None) -> dict:
    result = copy.deepcopy(base or {})
    if override:
        result.update(override)
    return result


def cm_or_none(value):
    return None if value is None else Cm(float(value))


def pt_or_none(value):
    return None if value is None else Pt(float(value))


def set_run_font(run, fmt: dict) -> None:
    font = run.font

    if "font_name" in fmt:
        font.name = fmt["font_name"]
        set_rfont(run, "w:ascii", fmt["font_name"])
        set_rfont(run, "w:hAnsi", fmt["font_name"])

    east_asia_font_name = fmt.get("east_asia_font_name") or fmt.get("font_name")
    if east_asia_font_name:
        set_rfont(run, "w:eastAsia", east_asia_font_name)

    if "font_size_pt" in fmt:
        font.size = Pt(float(fmt["font_size_pt"]))
    if "bold" in fmt:
        font.bold = bool(fmt["bold"])
    if "italic" in fmt:
        font.italic = bool(fmt["italic"])
    if "underline" in fmt:
        font.underline = bool(fmt["underline"])


def set_rfont(run, key: str, value: str) -> None:
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn(key), value)


def set_paragraph_format(paragraph, fmt: dict) -> None:
    paragraph_format = paragraph.paragraph_format

    if "alignment" in fmt:
        alignment = ALIGNMENTS.get(str(fmt["alignment"]).lower())
        if alignment is None:
            raise ValueError(f"不支持的段落对齐方式: {fmt['alignment']}")
        paragraph.alignment = alignment

    if "line_spacing" in fmt:
        paragraph_format.line_spacing = float(fmt["line_spacing"])
    if "first_line_indent_cm" in fmt:
        paragraph_format.first_line_indent = Cm(float(fmt["first_line_indent_cm"]))
    if "left_indent_cm" in fmt:
        paragraph_format.left_indent = cm_or_none(fmt["left_indent_cm"])
    if "right_indent_cm" in fmt:
        paragraph_format.right_indent = cm_or_none(fmt["right_indent_cm"])
    if "space_before_pt" in fmt:
        paragraph_format.space_before = pt_or_none(fmt["space_before_pt"])
    if "space_after_pt" in fmt:
        paragraph_format.space_after = pt_or_none(fmt["space_after_pt"])

    if not paragraph.runs:
        paragraph.add_run()

    for run in paragraph.runs:
        set_run_font(run, fmt)


def matches_rule(paragraph, index: int, rule: dict) -> bool:
    match = rule.get("match", {})
    text = paragraph.text.strip()

    if "paragraph_index" in match and int(match["paragraph_index"]) != index:
        return False

    if match.get("non_empty") is True and not text:
        return False

    if "starts_with" in match:
        prefixes = match["starts_with"]
        if isinstance(prefixes, str):
            prefixes = [prefixes]
        if not any(text.startswith(prefix) for prefix in prefixes):
            return False

    if "contains" in match and str(match["contains"]) not in text:
        return False

    if "regex" in match and not re.search(str(match["regex"]), text):
        return False

    return True


def choose_paragraph_format(paragraph, index: int, spec: dict) -> dict:
    fmt = copy.deepcopy(spec.get("default", {}))
    for rule in spec.get("rules", []):
        if matches_rule(paragraph, index, rule):
            fmt = merged_format(fmt, rule.get("format", {}))
            break
    return fmt


def apply_page_settings(document, spec: dict) -> None:
    page = spec.get("page", {})
    if not page:
        return

    for section in document.sections:
        orientation = str(page.get("orientation", "portrait")).lower()
        if orientation == "landscape":
            section.orientation = WD_ORIENT.LANDSCAPE
        elif orientation == "portrait":
            section.orientation = WD_ORIENT.PORTRAIT
        else:
            raise ValueError(f"不支持的页面方向: {page['orientation']}")

        size_name = str(page.get("size", "")).upper()
        if size_name in PAGE_SIZES_CM:
            width_cm, height_cm = PAGE_SIZES_CM[size_name]
            if orientation == "landscape":
                width_cm, height_cm = height_cm, width_cm
            section.page_width = Cm(width_cm)
            section.page_height = Cm(height_cm)

        margins = page.get("margins_cm", {})
        if "top" in margins:
            section.top_margin = Cm(float(margins["top"]))
        if "bottom" in margins:
            section.bottom_margin = Cm(float(margins["bottom"]))
        if "left" in margins:
            section.left_margin = Cm(float(margins["left"]))
        if "right" in margins:
            section.right_margin = Cm(float(margins["right"]))


def apply_table_settings(document, spec: dict) -> None:
    table_spec = spec.get("tables", {})
    if not table_spec or table_spec.get("enabled", True) is False:
        return

    base_fmt = {
        "font_name": table_spec.get("font_name", spec.get("default", {}).get("font_name")),
        "east_asia_font_name": table_spec.get(
            "east_asia_font_name", spec.get("default", {}).get("east_asia_font_name")
        ),
        "font_size_pt": table_spec.get("font_size_pt", spec.get("default", {}).get("font_size_pt")),
        "alignment": table_spec.get("alignment", "center"),
        "line_spacing": table_spec.get("line_spacing", 1.0),
        "first_line_indent_cm": table_spec.get("first_line_indent_cm", 0),
        "space_before_pt": table_spec.get("space_before_pt", 0),
        "space_after_pt": table_spec.get("space_after_pt", 0),
    }
    vertical_alignment = TABLE_VERTICAL_ALIGNMENTS.get(
        str(table_spec.get("cell_vertical_alignment", "center")).lower()
    )

    for table in document.tables:
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                if vertical_alignment is not None:
                    cell.vertical_alignment = vertical_alignment
                cell_fmt = copy.deepcopy(base_fmt)
                if row_index == 0 and table_spec.get("bold_header_row", False):
                    cell_fmt["bold"] = True
                for paragraph in cell.paragraphs:
                    set_paragraph_format(paragraph, cell_fmt)


def format_document(input_path: Path, output_path: Path, spec: dict) -> None:
    document = Document(str(input_path))
    apply_page_settings(document, spec)

    for index, paragraph in enumerate(document.paragraphs):
        fmt = choose_paragraph_format(paragraph, index, spec)
        set_paragraph_format(paragraph, fmt)

    apply_table_settings(document, spec)
    document.save(str(output_path))


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Format a .docx file according to the embedded specification.")
    parser.add_argument("input", type=Path, help="原始 .docx 文件")
    parser.add_argument("output", type=Path, help="输出 .docx 文件")
    parser.add_argument("--show-spec", action="store_true", help="打印当前脚本内置的格式规格")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    if args.show_spec:
        print(json.dumps(SPEC, ensure_ascii=False, indent=2))
        return

    if args.input.suffix.lower() != ".docx":
        raise SystemExit("当前工具只支持 .docx；旧版 .doc 需先转换成 .docx。")

    format_document(args.input, args.output, SPEC)
    print(f"已生成: {args.output}")


if __name__ == "__main__":
    main()
'''
    return dedent(template).replace("__SPEC_JSON__", repr(spec_literal))


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate a standalone Python Word formatter.")
    parser.add_argument("--spec", type=Path, required=True, help="JSON 格式要求文件")
    parser.add_argument("--output", type=Path, required=True, help="生成的 Python 脚本路径")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    spec = load_spec(args.spec)
    source = build_formatter_source(spec)
    args.output.parent.mkdir(parents=True, exist_ok=True)
    args.output.write_text(source, encoding="utf-8")
    args.output.chmod(0o755)
    print(f"已生成格式化脚本: {args.output}")


if __name__ == "__main__":
    main()
