#!/usr/bin/env python3
"""Validate formatted Word documents against format_spec.json."""

from __future__ import annotations

import argparse
import json
import re
import traceback
import zipfile
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from lxml import etree


UNKNOWN = "unknown"
CM_TOLERANCE = 0.05
PT_TOLERANCE = 0.5
LINE_SPACING_TOLERANCE = 0.05
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
HEADER_REL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/header"
FOOTER_REL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footer"

PAGE_SIZES_CM = {
    "A4": (21.0, 29.7),
    "LETTER": (21.59, 27.94),
}

ALIGNMENT_NAMES = {
    WD_ALIGN_PARAGRAPH.LEFT: "left",
    WD_ALIGN_PARAGRAPH.CENTER: "center",
    WD_ALIGN_PARAGRAPH.RIGHT: "right",
    WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
    None: None,
}

VERTICAL_ALIGNMENT_NAMES = {
    WD_ALIGN_VERTICAL.TOP: "top",
    WD_ALIGN_VERTICAL.CENTER: "center",
    WD_ALIGN_VERTICAL.BOTTOM: "bottom",
    None: None,
}


def is_known(value: Any) -> bool:
    return value not in (None, "", UNKNOWN)


def any_known(value: Any) -> bool:
    if isinstance(value, dict):
        return any(any_known(item) for item in value.values())
    if isinstance(value, list):
        return any(any_known(item) for item in value)
    return is_known(value)


def requested_patches(spec: dict[str, Any]) -> list[str]:
    raw = spec.get("openxml_patches", None)
    if raw is True:
        return ["captions", "equation_numbering", "header_footer", "math_font", "references", "three_line_table"]
    if raw is False:
        return []
    if raw is not None:
        if not raw:
            return []
        if isinstance(raw, list):
            return [str(item) for item in raw]
        return [str(raw)]

    patches = []
    tables = spec.get("tables", {}) if isinstance(spec.get("tables"), dict) else {}
    figures = spec.get("figures", {}) if isinstance(spec.get("figures"), dict) else {}
    equations = spec.get("equations", {}) if isinstance(spec.get("equations"), dict) else {}
    references = spec.get("references", {}) if isinstance(spec.get("references"), dict) else {}
    headers_footers = spec.get("headers_footers", {}) if isinstance(spec.get("headers_footers"), dict) else {}

    if tables.get("style") == "three_line":
        patches.append("three_line_table")
    if any_known(headers_footers):
        patches.append("header_footer")
    if any_known(figures) or is_known(tables.get("caption_position")) or is_known(tables.get("caption_alignment")):
        patches.append("captions")
    if any_known(references):
        patches.append("references")
    if is_known(equations.get("font")):
        patches.append("math_font")
    if is_known(equations.get("numbering")):
        patches.append("equation_numbering")
    return patches


def load_json(path: Path) -> dict[str, Any]:
    with path.open("r", encoding="utf-8") as handle:
        data = json.load(handle)
    if not isinstance(data, dict):
        raise ValueError(f"{path} 顶层必须是 JSON 对象")
    return data


def read_docx_xml(input_path: Path, member: str):
    try:
        with zipfile.ZipFile(input_path, "r") as package:
            return etree.fromstring(package.read(member))
    except KeyError:
        return None


def docx_members(input_path: Path) -> list[str]:
    with zipfile.ZipFile(input_path, "r") as package:
        return package.namelist()


def xml_text(root) -> str:
    if root is None:
        return ""
    return "".join(root.itertext())


def relationship_targets(input_path: Path, rel_type: str) -> list[str]:
    rels_root = read_docx_xml(input_path, "word/_rels/document.xml.rels")
    if rels_root is None:
        return []
    targets = []
    for rel in rels_root.findall(f"{{{REL_NS}}}Relationship"):
        if rel.get("Type") == rel_type and rel.get("Target"):
            target = rel.get("Target")
            targets.append(target if target.startswith("word/") else f"word/{target}")
    return targets


def read_related_xml_parts(input_path: Path, rel_type: str) -> list[tuple[str, Any]]:
    parts = []
    for target in relationship_targets(input_path, rel_type):
        root = read_docx_xml(input_path, target)
        if root is not None:
            parts.append((target, root))
    return parts


def cm_value(length: Any) -> float | None:
    if length is None:
        return None
    return round(float(length.cm), 4)


def pt_value(length: Any) -> float | None:
    if length is None:
        return None
    return round(float(length.pt), 4)


def normalize_alignment(value: Any) -> str | None:
    if value is None:
        return None
    text = str(value).lower()
    if text == "both":
        return "justify"
    return text


def first_text_run(paragraph):
    if not paragraph.runs:
        return None
    for run in paragraph.runs:
        if run.text.strip():
            return run
    return paragraph.runs[0]


def east_asia_font(run) -> str | None:
    if run is None or run._element.rPr is None or run._element.rPr.rFonts is None:
        return None
    return run._element.rPr.rFonts.get(qn("w:eastAsia"))


def ascii_font(run) -> str | None:
    if run is None:
        return None
    if run.font.name:
        return run.font.name
    if run._element.rPr is not None and run._element.rPr.rFonts is not None:
        return run._element.rPr.rFonts.get(qn("w:ascii")) or run._element.rPr.rFonts.get(qn("w:hAnsi"))
    return None


def run_font_size(run) -> float | None:
    if run is None or run.font.size is None:
        return None
    return pt_value(run.font.size)


def report_shell(input_path: Path, spec_path: Path, structure_path: Path) -> dict[str, Any]:
    return {
        "status": "pending",
        "input": str(input_path.resolve()),
        "spec": str(spec_path.resolve()),
        "structure": str(structure_path.resolve()),
        "summary": {"pass": 0, "warn": 0, "fail": 0},
        "checks": [],
        "warnings": [],
        "errors": [],
    }


def add_check(
    report: dict[str, Any],
    check_id: str,
    scope: str,
    target: str,
    expected: Any,
    actual: Any,
    result: str,
    message: str,
    unit: str | None = None,
) -> None:
    item = {
        "id": check_id,
        "scope": scope,
        "target": target,
        "expected": expected,
        "actual": actual,
        "result": result,
        "message": message,
    }
    if unit:
        item["unit"] = unit
    report["checks"].append(item)
    report["summary"][result] += 1
    if result == "warn":
        report["warnings"].append(message)


def compare_number(
    report: dict[str, Any],
    check_id: str,
    scope: str,
    target: str,
    expected: Any,
    actual: Any,
    tolerance: float,
    unit: str,
) -> None:
    if not is_known(expected):
        add_check(report, check_id, scope, target, expected, actual, "warn", f"{check_id} 规则未知，跳过校验", unit)
        return
    if actual is None:
        add_check(report, check_id, scope, target, expected, actual, "warn", f"{check_id} 无法读取实际值", unit)
        return
    expected_number = float(expected)
    result = "pass" if abs(float(actual) - expected_number) <= tolerance else "fail"
    message = f"{check_id} {'符合要求' if result == 'pass' else '不符合要求'}"
    add_check(report, check_id, scope, target, expected_number, actual, result, message, unit)


def compare_text(
    report: dict[str, Any],
    check_id: str,
    scope: str,
    target: str,
    expected: Any,
    actual: Any,
) -> None:
    if not is_known(expected):
        add_check(report, check_id, scope, target, expected, actual, "warn", f"{check_id} 规则未知，跳过校验")
        return
    if actual is None:
        add_check(report, check_id, scope, target, expected, actual, "warn", f"{check_id} 无法读取实际值")
        return
    expected_text = normalize_alignment(expected)
    actual_text = normalize_alignment(actual)
    result = "pass" if actual_text == expected_text else "fail"
    message = f"{check_id} {'符合要求' if result == 'pass' else '不符合要求'}"
    add_check(report, check_id, scope, target, expected_text, actual_text, result, message)


def compare_bool(
    report: dict[str, Any],
    check_id: str,
    scope: str,
    target: str,
    expected: Any,
    actual: Any,
) -> None:
    if not is_known(expected):
        add_check(report, check_id, scope, target, expected, actual, "warn", f"{check_id} 规则未知，跳过校验")
        return
    if actual is None:
        add_check(report, check_id, scope, target, expected, actual, "warn", f"{check_id} 无法读取实际值")
        return
    result = "pass" if bool(actual) == bool(expected) else "fail"
    message = f"{check_id} {'符合要求' if result == 'pass' else '不符合要求'}"
    add_check(report, check_id, scope, target, bool(expected), bool(actual), result, message)


def legacy_body_format(spec: dict[str, Any]) -> dict[str, Any]:
    return dict(spec.get("default", {}))


def module1_body_format(spec: dict[str, Any]) -> dict[str, Any]:
    body = spec.get("body", {})
    font = body.get("font", {}) if isinstance(body.get("font"), dict) else {}
    return {
        "font_name": font.get("latin"),
        "east_asia_font_name": font.get("east_asia"),
        "font_size_pt": font.get("size_pt"),
        "alignment": body.get("alignment"),
        "line_spacing": body.get("line_spacing"),
        "first_line_indent_chars": body.get("first_line_indent_chars"),
    }


def body_format(spec: dict[str, Any]) -> dict[str, Any]:
    return legacy_body_format(spec) if "default" in spec else module1_body_format(spec)


def merge_format(base: dict[str, Any], override: dict[str, Any]) -> dict[str, Any]:
    merged = dict(base)
    for key, value in override.items():
        if is_known(value):
            merged[key] = value
    return merged


def heading_format_from_module1(spec: dict[str, Any], level: int) -> dict[str, Any] | None:
    for rule in spec.get("headings", []):
        if int(rule.get("level", 0)) == int(level):
            font = rule.get("font", {}) if isinstance(rule.get("font"), dict) else {}
            return merge_format(
                body_format(spec),
                {
                    "font_name": font.get("latin"),
                    "east_asia_font_name": font.get("east_asia"),
                    "font_size_pt": font.get("size_pt"),
                    "alignment": rule.get("alignment"),
                    "bold": rule.get("bold"),
                    "first_line_indent_cm": 0,
                },
            )
    return None


def legacy_rule_matches(text: str, index: int, rule: dict[str, Any]) -> bool:
    match = rule.get("match", {})
    stripped = text.strip()
    if "paragraph_index" in match and int(match["paragraph_index"]) != index:
        return False
    if match.get("non_empty") is True and not stripped:
        return False
    if "starts_with" in match:
        prefixes = match["starts_with"]
        if isinstance(prefixes, str):
            prefixes = [prefixes]
        if not any(stripped.startswith(prefix) for prefix in prefixes):
            return False
    if "contains" in match and str(match["contains"]) not in stripped:
        return False
    if "regex" in match and not re.search(str(match["regex"]), stripped):
        return False
    return True


def paragraph_expected_format(spec: dict[str, Any], paragraph, index: int, block: dict[str, Any]) -> tuple[dict[str, Any] | None, str]:
    if "default" in spec:
        fmt = body_format(spec)
        for rule in spec.get("rules", []):
            if legacy_rule_matches(paragraph.text, index, rule):
                return merge_format(fmt, rule.get("format", {})), f"legacy.rule.{rule.get('name', index)}"
        return fmt, "legacy.default"

    if block.get("type") == "heading":
        fmt = heading_format_from_module1(spec, int(block.get("level", 0)))
        return fmt, f"heading.level_{block.get('level')}"
    if block.get("type") == "figure_caption":
        figures = spec.get("figures", {})
        font = figures.get("font", {}) if isinstance(figures.get("font"), dict) else {}
        return merge_format(
            body_format(spec),
            {
                "font_name": font.get("latin"),
                "east_asia_font_name": font.get("east_asia"),
                "font_size_pt": font.get("size_pt"),
                "alignment": figures.get("caption_alignment"),
                "first_line_indent_cm": 0,
            },
        ), "figure.caption"
    if block.get("type") == "table_caption":
        tables = spec.get("tables", {})
        font = tables.get("font", {}) if isinstance(tables.get("font"), dict) else {}
        return merge_format(
            body_format(spec),
            {
                "font_name": font.get("latin"),
                "east_asia_font_name": font.get("east_asia"),
                "font_size_pt": font.get("size_pt"),
                "alignment": tables.get("caption_alignment"),
                "first_line_indent_cm": 0,
            },
        ), "table.caption"
    return body_format(spec), "body"


def table_format(spec: dict[str, Any]) -> dict[str, Any]:
    tables = spec.get("tables", {})
    if "default" in spec:
        return {
            "font_name": tables.get("font_name", spec.get("default", {}).get("font_name")),
            "east_asia_font_name": tables.get("east_asia_font_name", spec.get("default", {}).get("east_asia_font_name")),
            "font_size_pt": tables.get("font_size_pt", spec.get("default", {}).get("font_size_pt")),
            "alignment": tables.get("alignment", "center"),
            "line_spacing": tables.get("line_spacing", 1.0),
            "first_line_indent_cm": tables.get("first_line_indent_cm", 0),
            "space_before_pt": tables.get("space_before_pt", 0),
            "space_after_pt": tables.get("space_after_pt", 0),
            "bold_header_row": tables.get("bold_header_row", False),
            "cell_vertical_alignment": tables.get("cell_vertical_alignment", "center"),
        }

    font = tables.get("font", {}) if isinstance(tables.get("font"), dict) else {}
    return {
        "font_name": font.get("latin"),
        "east_asia_font_name": font.get("east_asia"),
        "font_size_pt": font.get("size_pt"),
        "alignment": tables.get("caption_alignment", "center"),
        "line_spacing": 1.0,
        "first_line_indent_cm": 0,
        "bold_header_row": tables.get("bold_header_row", True),
        "cell_vertical_alignment": "center",
    }


def validate_page(document: Document, spec: dict[str, Any], report: dict[str, Any]) -> None:
    page = spec.get("page", {})
    if not isinstance(page, dict) or not page:
        add_check(report, "page", "page", "document", None, None, "warn", "未提供页面规则，跳过页面校验")
        return

    size_name = page.get("size", page.get("paper_size"))
    expected_size = PAGE_SIZES_CM.get(str(size_name).upper()) if is_known(size_name) else None
    expected_orientation = str(page.get("orientation")).lower() if is_known(page.get("orientation")) else None

    for section_index, section in enumerate(document.sections):
        target = f"section[{section_index}]"
        actual_orientation = "landscape" if section.orientation == WD_ORIENT.LANDSCAPE or section.page_width > section.page_height else "portrait"
        if expected_orientation:
            compare_text(report, "page.orientation", "page", target, expected_orientation, actual_orientation)

        if expected_size:
            width_cm, height_cm = expected_size
            if expected_orientation == "landscape":
                width_cm, height_cm = height_cm, width_cm
            compare_number(report, "page.width", "page", target, width_cm, cm_value(section.page_width), CM_TOLERANCE, "cm")
            compare_number(report, "page.height", "page", target, height_cm, cm_value(section.page_height), CM_TOLERANCE, "cm")
        elif is_known(size_name):
            add_check(report, "page.size", "page", target, size_name, None, "warn", f"暂不支持校验纸张类型 {size_name}")

        margins = page.get("margins_cm", {})
        if isinstance(margins, dict):
            for side, attr in (
                ("top", "top_margin"),
                ("bottom", "bottom_margin"),
                ("left", "left_margin"),
                ("right", "right_margin"),
            ):
                compare_number(
                    report,
                    f"page.margin.{side}",
                    "page",
                    target,
                    margins.get(side),
                    cm_value(getattr(section, attr)),
                    CM_TOLERANCE,
                    "cm",
                )

        if "header_distance_cm" in page:
            compare_number(
                report,
                "page.header_distance",
                "headers_footers",
                target,
                page.get("header_distance_cm"),
                cm_value(section.header_distance),
                CM_TOLERANCE,
                "cm",
            )
        if "footer_distance_cm" in page:
            compare_number(
                report,
                "page.footer_distance",
                "headers_footers",
                target,
                page.get("footer_distance_cm"),
                cm_value(section.footer_distance),
                CM_TOLERANCE,
                "cm",
            )


def validate_paragraph_format(report: dict[str, Any], paragraph, fmt: dict[str, Any], scope: str, target: str) -> None:
    paragraph_format = paragraph.paragraph_format
    run = first_text_run(paragraph)
    if "font_name" in fmt:
        compare_text(report, f"{scope}.font.latin", scope, target, fmt.get("font_name"), ascii_font(run))
    if "east_asia_font_name" in fmt:
        compare_text(report, f"{scope}.font.east_asia", scope, target, fmt.get("east_asia_font_name"), east_asia_font(run))
    if "font_size_pt" in fmt:
        compare_number(report, f"{scope}.font.size_pt", scope, target, fmt.get("font_size_pt"), run_font_size(run), PT_TOLERANCE, "pt")
    if "bold" in fmt:
        compare_bool(report, f"{scope}.bold", scope, target, fmt.get("bold"), run.font.bold if run is not None else None)
    if "alignment" in fmt:
        compare_text(report, f"{scope}.alignment", scope, target, fmt.get("alignment"), ALIGNMENT_NAMES.get(paragraph.alignment))
    if "line_spacing" in fmt:
        compare_number(
            report,
            f"{scope}.line_spacing",
            scope,
            target,
            fmt.get("line_spacing"),
            paragraph_format.line_spacing if isinstance(paragraph_format.line_spacing, (int, float)) else None,
            LINE_SPACING_TOLERANCE,
            "multiple",
        )

    if is_known(fmt.get("first_line_indent_cm")):
        expected_indent = fmt.get("first_line_indent_cm")
    elif is_known(fmt.get("first_line_indent_chars")):
        expected_indent = float(fmt["first_line_indent_chars"]) * 0.37
    else:
        expected_indent = None
    if "first_line_indent_cm" in fmt or "first_line_indent_chars" in fmt:
        compare_number(
            report,
            f"{scope}.first_line_indent",
            scope,
            target,
            expected_indent,
            cm_value(paragraph_format.first_line_indent),
            CM_TOLERANCE,
            "cm",
        )
    if "space_before_pt" in fmt:
        compare_number(report, f"{scope}.space_before", scope, target, fmt.get("space_before_pt"), pt_value(paragraph_format.space_before), PT_TOLERANCE, "pt")
    if "space_after_pt" in fmt:
        compare_number(report, f"{scope}.space_after", scope, target, fmt.get("space_after_pt"), pt_value(paragraph_format.space_after), PT_TOLERANCE, "pt")


def paragraph_blocks(structure: dict[str, Any]) -> list[dict[str, Any]]:
    supported = {"body", "heading", "figure_caption", "table_caption"}
    blocks = []
    for block in structure.get("blocks", []):
        source = block.get("source", {})
        if block.get("type") in supported and source.get("kind") == "paragraph":
            blocks.append(block)
    return blocks


def validate_paragraphs(document: Document, spec: dict[str, Any], structure: dict[str, Any], report: dict[str, Any]) -> None:
    blocks = paragraph_blocks(structure)
    if not blocks:
        add_check(report, "paragraphs", "paragraph", "document", None, None, "warn", "结构文件未提供可校验段落，跳过段落校验")
        return

    for block in blocks:
        source = block["source"]
        index = int(source["index"])
        target = f"paragraph[{index}]"
        if index >= len(document.paragraphs):
            add_check(report, "paragraph.index", "paragraph", target, index, len(document.paragraphs), "warn", f"{target} 超出文档段落范围")
            continue
        paragraph = document.paragraphs[index]
        fmt, scope = paragraph_expected_format(spec, paragraph, index, block)
        if fmt is None:
            add_check(report, f"{scope}.rule", scope, target, block.get("level"), None, "warn", f"{target} 未找到对应格式规则")
            continue
        validate_paragraph_format(report, paragraph, fmt, scope, target)

    if structure.get("preserve"):
        add_check(
            report,
            "preserve",
            "structure",
            "paper_structure.json",
            0,
            len(structure["preserve"]),
            "warn",
            f"结构文件包含 {len(structure['preserve'])} 个 preserve 项，模块5保留但不强判",
        )


def validate_tables(document: Document, spec: dict[str, Any], report: dict[str, Any]) -> None:
    tables_spec = spec.get("tables", {})
    if not isinstance(tables_spec, dict) or not tables_spec:
        add_check(report, "tables", "table", "document", None, len(document.tables), "warn", "未提供表格规则，跳过表格校验")
        return
    if not document.tables:
        add_check(report, "tables.exists", "table", "document", True, False, "warn", "格式规则包含表格要求，但文档中没有表格")
        return

    fmt = table_format(spec)
    for table_index, table in enumerate(document.tables):
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                target = f"table[{table_index}].row[{row_index}].cell[{cell_index}]"
                compare_text(
                    report,
                    "table.cell.vertical_alignment",
                    "table",
                    target,
                    fmt.get("cell_vertical_alignment"),
                    VERTICAL_ALIGNMENT_NAMES.get(cell.vertical_alignment),
                )
                if not cell.paragraphs:
                    add_check(report, "table.cell.paragraph", "table", target, True, False, "warn", f"{target} 没有可校验段落")
                    continue
                cell_scope = "table.header" if row_index == 0 and fmt.get("bold_header_row") else "table.cell"
                cell_fmt = dict(fmt)
                if row_index == 0 and fmt.get("bold_header_row"):
                    cell_fmt["bold"] = True
                validate_paragraph_format(report, cell.paragraphs[0], cell_fmt, cell_scope, target)


def xml_attr(element, attr: str) -> str | None:
    if element is None:
        return None
    return element.get(f"{{{W_NS}}}{attr}")


def border_value(root, path: str, attr: str) -> str | None:
    element = root.find(path)
    return xml_attr(element, attr)


def validate_three_line_table(input_path: Path, spec: dict[str, Any], report: dict[str, Any]) -> None:
    if "three_line_table" not in requested_patches(spec):
        return
    root = read_docx_xml(input_path, "word/document.xml")
    if root is None:
        add_check(report, "table.openxml.document_xml", "table.openxml", "word/document.xml", True, False, "warn", "无法读取 document.xml")
        return
    tables = root.findall(f".//{{{W_NS}}}tbl")
    if not tables:
        add_check(report, "table.openxml.exists", "table.openxml", "document", True, False, "warn", "需要校验三线表，但文档中没有表格")
        return

    for index, table in enumerate(tables):
        target = f"table[{index}]"
        borders = table.find(f"./{{{W_NS}}}tblPr/{{{W_NS}}}tblBorders")
        if borders is None:
            add_check(report, "table.openxml.borders", "table.openxml", target, True, False, "fail", f"{target} 缺少 tblBorders")
            continue
        compare_text(report, "table.openxml.top.val", "table.openxml", target, "single", border_value(borders, f"{{{W_NS}}}top", "val"))
        compare_text(report, "table.openxml.bottom.val", "table.openxml", target, "single", border_value(borders, f"{{{W_NS}}}bottom", "val"))
        compare_text(report, "table.openxml.top.sz", "table.openxml", target, "12", border_value(borders, f"{{{W_NS}}}top", "sz"))
        compare_text(report, "table.openxml.bottom.sz", "table.openxml", target, "12", border_value(borders, f"{{{W_NS}}}bottom", "sz"))
        for side in ("left", "right", "insideH", "insideV"):
            actual = border_value(borders, f"{{{W_NS}}}{side}", "val")
            result = "pass" if actual in (None, "nil", "none") else "fail"
            add_check(
                report,
                f"table.openxml.{side}.val",
                "table.openxml",
                target,
                "nil-or-missing",
                actual,
                result,
                f"table.openxml.{side}.val {'符合要求' if result == 'pass' else '不符合要求'}",
            )
        header_bottom = table.find(f"./{{{W_NS}}}tr[1]/{{{W_NS}}}tc/{{{W_NS}}}tcPr/{{{W_NS}}}tcBorders/{{{W_NS}}}bottom")
        compare_text(report, "table.openxml.header.bottom.val", "table.openxml", target, "single", xml_attr(header_bottom, "val"))
        compare_text(report, "table.openxml.header.bottom.sz", "table.openxml", target, "6", xml_attr(header_bottom, "sz"))


def validate_headers_footers(input_path: Path, spec: dict[str, Any], report: dict[str, Any]) -> None:
    headers_footers = spec.get("headers_footers", {})
    if not isinstance(headers_footers, dict) or not headers_footers:
        return
    header = headers_footers.get("header", {}) if isinstance(headers_footers.get("header"), dict) else {}
    footer = headers_footers.get("footer", {}) if isinstance(headers_footers.get("footer"), dict) else {}
    document_root = read_docx_xml(input_path, "word/document.xml")
    if document_root is not None:
        add_check(report, "headers_footers.header_reference", "headers_footers.openxml", "document", True, document_root.find(f".//{{{W_NS}}}headerReference") is not None, "pass" if document_root.find(f".//{{{W_NS}}}headerReference") is not None or not any_known(header) else "fail", "页眉引用检查")
        add_check(report, "headers_footers.footer_reference", "headers_footers.openxml", "document", True, document_root.find(f".//{{{W_NS}}}footerReference") is not None, "pass" if document_root.find(f".//{{{W_NS}}}footerReference") is not None or not any_known(footer) else "fail", "页脚引用检查")

    header_parts = read_related_xml_parts(input_path, HEADER_REL)
    footer_parts = read_related_xml_parts(input_path, FOOTER_REL)
    if any_known(header):
        add_check(report, "headers_footers.header_relationship", "headers_footers.openxml", "document.xml.rels", True, bool(header_parts), "pass" if header_parts else "fail", "页眉 relationship 检查")
    if any_known(footer):
        add_check(report, "headers_footers.footer_relationship", "headers_footers.openxml", "document.xml.rels", True, bool(footer_parts), "pass" if footer_parts else "fail", "页脚 relationship 检查")

    if is_known(header.get("text")):
        actual_text = "\n".join(xml_text(root) for _, root in header_parts)
        result = "pass" if str(header["text"]) in actual_text else "fail"
        add_check(report, "headers_footers.header.text", "headers_footers.openxml", "header", header["text"], actual_text, result, "页眉文本检查")

    page_number = footer.get("page_number")
    if page_number is True or str(page_number).lower() in {"true", "page", "page_number"}:
        has_page = any("PAGE" in xml_text(root) for _, root in footer_parts)
        add_check(report, "headers_footers.footer.page_number", "headers_footers.openxml", "footer", True, has_page, "pass" if has_page else "fail", "页脚页码字段检查")


def validate_caption_patch(document: Document, spec: dict[str, Any], structure: dict[str, Any], report: dict[str, Any]) -> None:
    if "captions" not in requested_patches(spec):
        return
    caption_blocks = [block for block in paragraph_blocks(structure) if block.get("type") in {"figure_caption", "table_caption"}]
    if not caption_blocks:
        add_check(report, "caption.exists", "caption", "paper_structure.json", True, False, "warn", "需要校验图表题，但结构文件没有图表题段落")
        return
    for block in caption_blocks:
        index = int(block["source"]["index"])
        if index >= len(document.paragraphs):
            add_check(report, "caption.index", "caption", f"paragraph[{index}]", index, len(document.paragraphs), "warn", "图表题段落索引超出范围")
            continue
        paragraph = document.paragraphs[index]
        fmt, scope = paragraph_expected_format(spec, paragraph, index, block)
        validate_paragraph_format(report, paragraph, fmt or {}, scope, f"paragraph[{index}]")
        compare_number(report, f"{scope}.first_line_indent", scope, f"paragraph[{index}]", 0, cm_value(paragraph.paragraph_format.first_line_indent), CM_TOLERANCE, "cm")
        if paragraph.paragraph_format.keep_with_next is None:
            add_check(report, f"{scope}.keep_with_next", scope, f"paragraph[{index}]", True, None, "warn", "无法可靠读取图表题 keep_with_next")
        else:
            compare_bool(report, f"{scope}.keep_with_next", scope, f"paragraph[{index}]", True, paragraph.paragraph_format.keep_with_next)


def reference_blocks(document: Document, structure: dict[str, Any]) -> list[int]:
    indexes = []
    for block in structure.get("references", []):
        source = block.get("source", {})
        if source.get("kind") == "paragraph":
            indexes.append(int(source["index"]))
    if indexes:
        return indexes

    in_references = False
    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        normalized = re.sub(r"[\s:：]+", "", text).lower()
        if normalized in {"参考文献", "references"}:
            in_references = True
            continue
        if in_references and normalized in {"致谢", "附录", "acknowledgements", "appendix"}:
            break
        if in_references and text:
            indexes.append(index)
    return indexes


def validate_references_patch(document: Document, spec: dict[str, Any], structure: dict[str, Any], report: dict[str, Any]) -> None:
    if "references" not in requested_patches(spec):
        return
    references = spec.get("references", {}) if isinstance(spec.get("references"), dict) else {}
    indexes = reference_blocks(document, structure)
    if not indexes:
        add_check(report, "references.exists", "references", "document", True, False, "warn", "需要校验参考文献，但未定位到参考文献条目")
        return
    for index in indexes:
        paragraph = document.paragraphs[index]
        target = f"paragraph[{index}]"
        compare_text(report, "references.alignment", "references", target, references.get("alignment", "justify"), ALIGNMENT_NAMES.get(paragraph.alignment))
        if references.get("indent") == "hanging":
            compare_number(report, "references.left_indent", "references", target, 0.74, cm_value(paragraph.paragraph_format.left_indent), CM_TOLERANCE, "cm")
            compare_number(report, "references.first_line_indent", "references", target, -0.74, cm_value(paragraph.paragraph_format.first_line_indent), CM_TOLERANCE, "cm")
        elif is_known(references.get("indent")):
            compare_number(report, "references.first_line_indent", "references", target, 0, cm_value(paragraph.paragraph_format.first_line_indent), CM_TOLERANCE, "cm")
        if "\xa0" in paragraph.text:
            add_check(report, "references.nbsp", "references", target, False, True, "fail", "参考文献中仍存在 NBSP")
        else:
            add_check(report, "references.nbsp", "references", target, False, False, "pass", "参考文献 NBSP 检查通过")


def validate_math_font_patch(input_path: Path, spec: dict[str, Any], report: dict[str, Any]) -> None:
    if "math_font" not in requested_patches(spec):
        return
    equations = spec.get("equations", {}) if isinstance(spec.get("equations"), dict) else {}
    expected_font = equations.get("font", "Times New Roman")
    settings_root = read_docx_xml(input_path, "word/settings.xml")
    math_font = settings_root.find(f".//{{{M_NS}}}mathFont") if settings_root is not None else None
    compare_text(report, "equations.math_font.settings", "equations.math_font", "word/settings.xml", expected_font, math_font.get(f"{{{M_NS}}}val") if math_font is not None else None)

    document_root = read_docx_xml(input_path, "word/document.xml")
    math_runs = list(document_root.iter(f"{{{M_NS}}}r")) if document_root is not None else []
    if not math_runs:
        add_check(report, "equations.math_font.runs", "equations.math_font", "word/document.xml", expected_font, 0, "warn", "文档中没有可校验的数学 run")
        return
    for index, math_run in enumerate(math_runs[:20]):
        fonts = math_run.find(f".//{{{W_NS}}}rFonts")
        actual = fonts.get(f"{{{W_NS}}}ascii") if fonts is not None else None
        compare_text(report, "equations.math_font.run", "equations.math_font", f"math_run[{index}]", expected_font, actual)


def validate_equation_numbering_patch(document: Document, spec: dict[str, Any], report: dict[str, Any]) -> None:
    if "equation_numbering" not in requested_patches(spec):
        return
    equations = spec.get("equations", {}) if isinstance(spec.get("equations"), dict) else {}
    numbering = equations.get("numbering")
    equation_indexes = []
    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        has_math = paragraph._element.find(f".//{{{M_NS}}}oMath") is not None or paragraph._element.find(f".//{{{M_NS}}}oMathPara") is not None
        has_number = bool(re.search(r"[（(]\s*\d+(?:[-.]\d+)*\s*[）)]$", text))
        if has_math or has_number:
            equation_indexes.append(index)
    if not equation_indexes:
        add_check(report, "equations.numbering.exists", "equations.numbering", "document", numbering, None, "warn", "需要校验公式编号，但未定位到公式段落")
        return
    for index in equation_indexes:
        compare_text(report, "equations.numbering.alignment", "equations.numbering", f"paragraph[{index}]", "center", ALIGNMENT_NAMES.get(document.paragraphs[index].alignment))
    if numbering == "right_aligned":
        add_check(report, "equations.numbering.right_aligned", "equations.numbering", "document", numbering, "centered-paragraph-only", "warn", "当前模块2/5仅校验保守居中实现，未强校验右侧编号 eqArr")


def finalize_status(report: dict[str, Any]) -> None:
    if report["errors"]:
        report["status"] = "failed"
    elif report["summary"]["fail"]:
        report["status"] = "fail"
    elif report["summary"]["warn"]:
        report["status"] = "warn"
    else:
        report["status"] = "pass"


def validate_document(input_path: Path, spec_path: Path, structure_path: Path, report_json: Path | None = None, report_md: Path | None = None) -> dict[str, Any]:
    report = report_shell(input_path, spec_path, structure_path)
    try:
        if input_path.suffix.lower() != ".docx":
            raise ValueError("当前只支持 .docx")
        if not input_path.is_file():
            raise FileNotFoundError(f"输入文件不存在: {input_path}")
        if not spec_path.is_file():
            raise FileNotFoundError(f"format_spec.json 不存在: {spec_path}")
        if not structure_path.is_file():
            raise FileNotFoundError(f"paper_structure.json 不存在: {structure_path}")

        spec = load_json(spec_path)
        structure = load_json(structure_path)
        document = Document(str(input_path))
        validate_page(document, spec, report)
        validate_paragraphs(document, spec, structure, report)
        validate_tables(document, spec, report)
        validate_three_line_table(input_path, spec, report)
        validate_headers_footers(input_path, spec, report)
        validate_caption_patch(document, spec, structure, report)
        validate_references_patch(document, spec, structure, report)
        validate_math_font_patch(input_path, spec, report)
        validate_equation_numbering_patch(document, spec, report)
    except Exception as exc:
        report["errors"].append({"message": str(exc), "traceback": traceback.format_exc()})

    finalize_status(report)
    if report_json is not None:
        write_json_report(report_json, report)
    if report_md is not None:
        write_markdown_report(report_md, report)
    return report


def write_json_report(path: Path, report: dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")


def render_markdown_report(report: dict[str, Any]) -> str:
    lines = [
        "# 格式合规性校验报告",
        "",
        f"- 状态：{report['status']}",
        f"- 输入：{report['input']}",
        f"- 规格：{report['spec']}",
        f"- 结构：{report['structure']}",
        f"- 通过：{report['summary']['pass']}",
        f"- 警告：{report['summary']['warn']}",
        f"- 失败：{report['summary']['fail']}",
        "",
    ]

    for group in ("page", "body", "heading", "figure", "caption", "table", "references", "headers_footers", "equations", "structure", "paragraph"):
        group_checks = [item for item in report["checks"] if item["scope"].startswith(group) or item["id"].startswith(group)]
        if not group_checks:
            continue
        lines.extend([f"## {group}", ""])
        for item in group_checks:
            unit = f" {item['unit']}" if item.get("unit") else ""
            lines.append(
                f"- [{item['result']}] {item['id']} {item['target']}："
                f" expected={item['expected']} actual={item['actual']}{unit}"
            )
        lines.append("")

    if report["errors"]:
        lines.extend(["## errors", ""])
        for error in report["errors"]:
            lines.append(f"- {error.get('message', error)}")
        lines.append("")
    return "\n".join(lines)


def write_markdown_report(path: Path, report: dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(render_markdown_report(report), encoding="utf-8")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Validate a formatted .docx against format_spec.json.")
    subparsers = parser.add_subparsers(dest="command", required=True)

    validate = subparsers.add_parser("validate", help="validate formatted .docx")
    validate.add_argument("--input", type=Path, required=True, help="formatted.docx")
    validate.add_argument("--spec", type=Path, required=True, help="format_spec.json")
    validate.add_argument("--structure", type=Path, required=True, help="paper_structure.json")
    validate.add_argument("--report-json", type=Path, required=True, help="validation_report.json")
    validate.add_argument("--report-md", type=Path, required=True, help="validation_report.md")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    if args.command == "validate":
        report = validate_document(args.input, args.spec, args.structure, args.report_json, args.report_md)
        print(f"已生成: {args.report_json}")
        print(f"已生成: {args.report_md}")
        if report["status"] in {"failed", "fail"}:
            raise SystemExit(1)
        return
    raise RuntimeError(f"未知命令: {args.command}")


if __name__ == "__main__":
    main()
