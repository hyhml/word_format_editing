#!/usr/bin/env python3
"""Parse format requirement files into human and machine readable specs."""

from __future__ import annotations

import argparse
import copy
import json
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Iterable


UNKNOWN = "unknown"

CN_FONT_SIZES_PT = {
    "初号": 42,
    "小初": 36,
    "一号": 26,
    "小一": 24,
    "二号": 22,
    "小二": 18,
    "三号": 16,
    "小三": 15,
    "四号": 14,
    "小四": 12,
    "五号": 10.5,
    "小五": 9,
    "六号": 7.5,
    "小六": 6.5,
    "七号": 5.5,
    "八号": 5,
}
CN_SIZE_NAMES_BY_LENGTH = sorted(CN_FONT_SIZES_PT, key=len, reverse=True)

EAST_ASIA_FONTS = ["宋体", "黑体", "楷体", "仿宋", "隶书", "微软雅黑"]
LATIN_FONTS = ["Times New Roman", "Arial", "Calibri"]


class FormatParserError(ValueError):
    """Raised when a format requirement source cannot be parsed."""


@dataclass
class ExtractedSource:
    path: str
    type: str
    text: str
    paragraphs: list[str] = field(default_factory=list)
    tables: list[list[list[str]]] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)


@dataclass
class FieldCandidate:
    field: str
    value: Any
    source: str
    confidence: float
    source_text: str


def empty_spec() -> dict[str, Any]:
    return {
        "metadata": {
            "name": UNKNOWN,
            "institution": UNKNOWN,
            "document_type": UNKNOWN,
            "year": UNKNOWN,
            "description": "",
            "sources": [],
        },
        "page": {
            "paper_size": UNKNOWN,
            "orientation": UNKNOWN,
            "margins_cm": {
                "top": UNKNOWN,
                "bottom": UNKNOWN,
                "left": UNKNOWN,
                "right": UNKNOWN,
            },
            "header_distance_cm": UNKNOWN,
            "footer_distance_cm": UNKNOWN,
        },
        "body": {
            "font": {
                "east_asia": UNKNOWN,
                "latin": UNKNOWN,
                "size_cn": UNKNOWN,
                "size_pt": UNKNOWN,
            },
            "alignment": UNKNOWN,
            "line_spacing": UNKNOWN,
            "first_line_indent_chars": UNKNOWN,
        },
        "headings": [],
        "tables": {
            "style": UNKNOWN,
            "font": {
                "east_asia": UNKNOWN,
                "latin": UNKNOWN,
                "size_cn": UNKNOWN,
                "size_pt": UNKNOWN,
            },
            "caption_position": UNKNOWN,
            "caption_alignment": UNKNOWN,
        },
        "figures": {
            "caption_position": UNKNOWN,
            "caption_alignment": UNKNOWN,
            "font": {
                "east_asia": UNKNOWN,
                "latin": UNKNOWN,
                "size_cn": UNKNOWN,
                "size_pt": UNKNOWN,
            },
        },
        "equations": {
            "font": UNKNOWN,
            "alignment": UNKNOWN,
            "numbering": UNKNOWN,
        },
        "references": {
            "style": UNKNOWN,
            "alignment": UNKNOWN,
            "indent": UNKNOWN,
        },
        "headers_footers": {
            "header": {
                "text": UNKNOWN,
                "font": UNKNOWN,
                "size_cn": UNKNOWN,
                "size_pt": UNKNOWN,
                "alignment": UNKNOWN,
            },
            "footer": {
                "page_number": UNKNOWN,
                "font": UNKNOWN,
                "size_cn": UNKNOWN,
                "size_pt": UNKNOWN,
                "alignment": UNKNOWN,
            },
        },
        "derived_rules": [],
        "conflicts": [],
        "unknowns": [],
        "validation_rules": [],
    }


def read_text_file(path: Path) -> tuple[str, list[str]]:
    warnings = []
    for encoding in ("utf-8", "gb18030", "latin-1"):
        try:
            return path.read_text(encoding=encoding), warnings
        except UnicodeDecodeError:
            warnings.append(f"{path.name}: 使用 {encoding} 解码失败")
    raise FormatParserError(f"无法读取文本文件: {path}")


def extract_txt(path: Path) -> ExtractedSource:
    text, warnings = read_text_file(path)
    paragraphs = [line.strip() for line in text.splitlines() if line.strip()]
    return ExtractedSource(
        path=str(path),
        type=path.suffix.lower().lstrip(".") or "txt",
        text=text,
        paragraphs=paragraphs,
        warnings=warnings,
    )


def extract_docx(path: Path) -> ExtractedSource:
    try:
        from docx import Document
    except ImportError as exc:
        raise FormatParserError("解析 .docx 需要安装 python-docx") from exc

    document = Document(str(path))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    tables = []
    table_texts = []
    for table in document.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
            table_texts.append(" | ".join(cell for cell in cells if cell))
        tables.append(rows)
    text = "\n".join(paragraphs + [line for line in table_texts if line])
    return ExtractedSource(path=str(path), type="docx", text=text, paragraphs=paragraphs, tables=tables)


def extract_pdf(path: Path) -> ExtractedSource:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise FormatParserError("解析 .pdf 需要安装 pypdf") from exc

    warnings = []
    reader = PdfReader(str(path))
    pages = []
    for index, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if not text.strip():
            warnings.append(f"{path.name}: 第 {index + 1} 页未提取到文本，可能是扫描版 PDF")
        pages.append(text)
    text = "\n".join(pages)
    paragraphs = [line.strip() for line in text.splitlines() if line.strip()]
    if not text.strip():
        warnings.append(f"{path.name}: PDF 文本为空，当前版本暂不执行 OCR")
    return ExtractedSource(path=str(path), type="pdf", text=text, paragraphs=paragraphs, warnings=warnings)


def extract_source(path: Path) -> ExtractedSource:
    if not path.is_file():
        raise FormatParserError(f"格式要求文件不存在或不是文件: {path}")
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md", ".json"}:
        return extract_txt(path)
    if suffix == ".docx":
        return extract_docx(path)
    if suffix == ".pdf":
        return extract_pdf(path)
    raise FormatParserError(f"暂不支持的格式要求文件类型: {path.suffix}")


def set_nested(target: dict[str, Any], path: str, value: Any) -> None:
    keys = path.split(".")
    current = target
    for key in keys[:-1]:
        current = current[key]
    current[keys[-1]] = value


def get_nested(target: dict[str, Any], path: str) -> Any:
    current: Any = target
    for key in path.split("."):
        current = current[key]
    return current


def normalize_text(text: str) -> str:
    return re.sub(r"\s+", " ", text.replace("\u3000", " ")).strip()


def context_snippet(text: str, start: int, end: int, window: int = 40) -> str:
    left = max(0, start - window)
    right = min(len(text), end + window)
    return text[left:right].strip()


def add_candidate(
    candidates: dict[str, list[FieldCandidate]],
    field: str,
    value: Any,
    source: ExtractedSource,
    confidence: float,
    source_text: str,
) -> None:
    if value in (None, "", UNKNOWN):
        return
    candidates.setdefault(field, []).append(
        FieldCandidate(
            field=field,
            value=value,
            source=source.path,
            confidence=confidence,
            source_text=source_text,
        )
    )


def cn_size_to_pt(size_name: str) -> float | int | str:
    return CN_FONT_SIZES_PT.get(size_name, UNKNOWN)


def find_first_cn_size(text: str) -> str | None:
    matches = []
    for size_name in CN_SIZE_NAMES_BY_LENGTH:
        index = text.find(size_name)
        if index >= 0:
            matches.append((index, -len(size_name), size_name))
    if not matches:
        return None
    return sorted(matches)[0][2]


def collect_metadata_candidates(
    source: ExtractedSource,
    description: str,
    candidates: dict[str, list[FieldCandidate]],
) -> None:
    text = f"{description}\n{source.text}\n{Path(source.path).stem}"
    normalized = normalize_text(text)
    year = re.search(r"(20\d{2})", normalized)
    if year:
        add_candidate(candidates, "metadata.year", year.group(1), source, 0.8, year.group(0))

    institution = re.search(r"([\u4e00-\u9fff]{2,30}(?:大学|学院|学校|研究院|期刊|出版社))", normalized)
    if institution:
        add_candidate(candidates, "metadata.institution", institution.group(1), source, 0.75, institution.group(0))

    if "本科" in normalized and "论文" in normalized:
        add_candidate(candidates, "metadata.document_type", "本科毕业论文", source, 0.75, "本科...论文")
    elif "硕士" in normalized and "论文" in normalized:
        add_candidate(candidates, "metadata.document_type", "硕士论文", source, 0.75, "硕士...论文")
    elif "期刊" in normalized or "投稿" in normalized:
        add_candidate(candidates, "metadata.document_type", "期刊论文", source, 0.65, "期刊/投稿")


def collect_page_candidates(source: ExtractedSource, candidates: dict[str, list[FieldCandidate]]) -> None:
    text = normalize_text(source.text)
    if re.search(r"A\s*4|Ａ\s*4|A\s*４", text, re.IGNORECASE):
        add_candidate(candidates, "page.paper_size", "A4", source, 0.95, "A4")
    if re.search(r"横向|landscape", text, re.IGNORECASE):
        add_candidate(candidates, "page.orientation", "landscape", source, 0.7, "横向/landscape")
    if re.search(r"纵向|portrait", text, re.IGNORECASE):
        add_candidate(candidates, "page.orientation", "portrait", source, 0.7, "纵向/portrait")

    shared_margin = re.search(r"页边距[^。；;\n]{0,30}?(\d+(?:\.\d+)?)\s*(?:cm|厘米)", text)
    if shared_margin:
        value = float(shared_margin.group(1))
        snippet = context_snippet(text, shared_margin.start(), shared_margin.end())
        for side in ("top", "bottom", "left", "right"):
            add_candidate(candidates, f"page.margins_cm.{side}", value, source, 0.65, snippet)

    side_patterns = {
        "top": r"(?:上|顶部|上边距)[^0-9]{0,12}(\d+(?:\.\d+)?)\s*(?:cm|厘米)",
        "bottom": r"(?:下|底部|下边距)[^0-9]{0,12}(\d+(?:\.\d+)?)\s*(?:cm|厘米)",
        "left": r"(?:左|左边距)[^0-9]{0,12}(\d+(?:\.\d+)?)\s*(?:cm|厘米)",
        "right": r"(?:右|右边距)[^0-9]{0,12}(\d+(?:\.\d+)?)\s*(?:cm|厘米)",
    }
    for side, pattern in side_patterns.items():
        match = re.search(pattern, text)
        if match:
            add_candidate(
                candidates,
                f"page.margins_cm.{side}",
                float(match.group(1)),
                source,
                0.85,
                context_snippet(text, match.start(), match.end()),
            )

    header = re.search(r"页眉[^0-9]{0,12}(\d+(?:\.\d+)?)\s*(?:cm|厘米)", text)
    if header:
        add_candidate(candidates, "page.header_distance_cm", float(header.group(1)), source, 0.7, header.group(0))
    footer = re.search(r"页脚[^0-9]{0,12}(\d+(?:\.\d+)?)\s*(?:cm|厘米)", text)
    if footer:
        add_candidate(candidates, "page.footer_distance_cm", float(footer.group(1)), source, 0.7, footer.group(0))


def collect_body_candidates(source: ExtractedSource, candidates: dict[str, list[FieldCandidate]]) -> None:
    text = normalize_text(source.text)
    sentences = [part.strip() for part in re.split(r"[。；;\n]", source.text) if part.strip()]
    body_sentences = [sentence for sentence in sentences if re.search(r"正文|主体|段落", sentence)]
    search_text = normalize_text(" ".join(body_sentences)) if body_sentences else text

    for font in EAST_ASIA_FONTS:
        if font in search_text:
            add_candidate(candidates, "body.font.east_asia", font, source, 0.75, font)
            break
    for font in LATIN_FONTS:
        if font.lower() in search_text.lower():
            add_candidate(candidates, "body.font.latin", font, source, 0.75, font)
            break

    size_name = find_first_cn_size(search_text)
    if size_name:
        add_candidate(candidates, "body.font.size_cn", size_name, source, 0.75, size_name)
        add_candidate(candidates, "body.font.size_pt", cn_size_to_pt(size_name), source, 0.75, size_name)

    line_spacing = re.search(r"(\d+(?:\.\d+)?)\s*倍行距", search_text)
    if line_spacing:
        add_candidate(candidates, "body.line_spacing", float(line_spacing.group(1)), source, 0.85, line_spacing.group(0))

    indent = re.search(r"首行缩进[^0-9一二两两字符]{0,8}([0-9]+|一|二|两)\s*(?:个)?字符", search_text)
    if indent:
        value_text = indent.group(1)
        value = {"一": 1, "二": 2, "两": 2}.get(value_text, int(value_text) if value_text.isdigit() else UNKNOWN)
        add_candidate(candidates, "body.first_line_indent_chars", value, source, 0.85, indent.group(0))

    if "两端对齐" in search_text:
        add_candidate(candidates, "body.alignment", "justify", source, 0.85, "两端对齐")
    elif "居中" in search_text:
        add_candidate(candidates, "body.alignment", "center", source, 0.6, "居中")
    elif "左对齐" in search_text:
        add_candidate(candidates, "body.alignment", "left", source, 0.6, "左对齐")


def collect_heading_candidates(source: ExtractedSource) -> list[dict[str, Any]]:
    text = normalize_text(source.text)
    headings: dict[int, dict[str, Any]] = {}
    patterns = [
        (1, r"(?:一级标题|一[、级]标题|章标题)[^。；;\n]{0,60}"),
        (2, r"(?:二级标题|二[、级]标题|节标题)[^。；;\n]{0,60}"),
        (3, r"(?:三级标题|三[、级]标题)[^。；;\n]{0,60}"),
    ]
    for level, pattern in patterns:
        for match in re.finditer(pattern, text):
            snippet = match.group(0)
            rule = headings.setdefault(
                level,
                {
                    "level": level,
                    "font": {
                        "east_asia": UNKNOWN,
                        "latin": UNKNOWN,
                        "size_cn": UNKNOWN,
                        "size_pt": UNKNOWN,
                    },
                    "alignment": UNKNOWN,
                    "bold": UNKNOWN,
                    "source": source.path,
                    "source_text": snippet,
                },
            )
            for font in EAST_ASIA_FONTS:
                if font in snippet:
                    rule["font"]["east_asia"] = font
                    break
            for font in LATIN_FONTS:
                if font.lower() in snippet.lower():
                    rule["font"]["latin"] = font
                    break
            size_name = find_first_cn_size(snippet)
            if size_name:
                rule["font"]["size_cn"] = size_name
                rule["font"]["size_pt"] = cn_size_to_pt(size_name)
            if "居中" in snippet:
                rule["alignment"] = "center"
            elif "左对齐" in snippet or "顶格" in snippet:
                rule["alignment"] = "left"
            if "加粗" in snippet or "粗体" in snippet:
                rule["bold"] = True
            elif "不加粗" in snippet:
                rule["bold"] = False
    return [headings[level] for level in sorted(headings)]


def collect_special_candidates(source: ExtractedSource, candidates: dict[str, list[FieldCandidate]]) -> list[dict[str, Any]]:
    text = normalize_text(source.text)
    derived_rules = []

    if "三线表" in text or "三线制" in text:
        add_candidate(candidates, "tables.style", "three_line", source, 0.9, "三线表/三线制")
    if re.search(r"表题[^。；;\n]{0,30}(?:表上|上方|表前)", text):
        add_candidate(candidates, "tables.caption_position", "above", source, 0.75, "表题在表上方")
    if re.search(r"图题[^。；;\n]{0,30}(?:图下|下方|图后)", text):
        add_candidate(candidates, "figures.caption_position", "below", source, 0.75, "图题在图下方")
    if "图题" in text and "居中" in text:
        add_candidate(candidates, "figures.caption_alignment", "center", source, 0.65, "图题...居中")
    if "表题" in text and "居中" in text:
        add_candidate(candidates, "tables.caption_alignment", "center", source, 0.65, "表题...居中")

    caption_window = " ".join(
        context_snippet(text, match.start(), match.end(), 80)
        for match in re.finditer(r"图题|表题|图表标题|图表题", text)
    )
    for font in EAST_ASIA_FONTS:
        if font in caption_window:
            add_candidate(candidates, "figures.font.east_asia", font, source, 0.65, font)
            add_candidate(candidates, "tables.font.east_asia", font, source, 0.65, font)
            break
    size_name = find_first_cn_size(caption_window)
    if size_name:
        add_candidate(candidates, "figures.font.size_cn", size_name, source, 0.65, size_name)
        add_candidate(candidates, "figures.font.size_pt", cn_size_to_pt(size_name), source, 0.65, size_name)
        add_candidate(candidates, "tables.font.size_cn", size_name, source, 0.65, size_name)
        add_candidate(candidates, "tables.font.size_pt", cn_size_to_pt(size_name), source, 0.65, size_name)

    if "公式" in text and "居中" in text:
        add_candidate(candidates, "equations.alignment", "center", source, 0.65, "公式...居中")
    if "公式" in text and re.search(r"右[侧对齐]*编号|编号[^。；;\n]{0,20}右", text):
        add_candidate(candidates, "equations.numbering", "right_aligned", source, 0.7, "公式编号右对齐")
    if "Times New Roman" in text and "公式" in text:
        add_candidate(candidates, "equations.font", "Times New Roman", source, 0.65, "公式...Times New Roman")

    if "参考文献" in text:
        if "左对齐" in text:
            add_candidate(candidates, "references.alignment", "left", source, 0.65, "参考文献...左对齐")
        if "两端对齐" in text:
            add_candidate(candidates, "references.alignment", "justify", source, 0.65, "参考文献...两端对齐")
        if "悬挂缩进" in text:
            add_candidate(candidates, "references.indent", "hanging", source, 0.65, "参考文献...悬挂缩进")

    header_match = re.search(r"页眉[^。；;\n]{0,80}", text)
    if header_match:
        snippet = header_match.group(0)
        if "居中" in snippet:
            add_candidate(candidates, "headers_footers.header.alignment", "center", source, 0.65, snippet)
        for font in EAST_ASIA_FONTS:
            if font in snippet:
                add_candidate(candidates, "headers_footers.header.font", font, source, 0.65, snippet)
                break
        size_name = find_first_cn_size(snippet)
        if size_name:
            add_candidate(candidates, "headers_footers.header.size_cn", size_name, source, 0.65, snippet)
            add_candidate(candidates, "headers_footers.header.size_pt", cn_size_to_pt(size_name), source, 0.65, snippet)

    footer_match = re.search(r"页脚|页码", text)
    if footer_match:
        add_candidate(candidates, "headers_footers.footer.page_number", True, source, 0.55, footer_match.group(0))
        if "居中" in text:
            add_candidate(candidates, "headers_footers.footer.alignment", "center", source, 0.55, "页脚/页码...居中")

    for match in re.finditer(r"(?:比正文小一号|较正文小一号)", text):
        derived_rules.append(
            {
                "type": "relative_font_size",
                "target": "caption",
                "base": "body",
                "operation": "smaller_by",
                "steps": 1,
                "source": source.path,
                "source_text": context_snippet(text, match.start(), match.end()),
            }
        )
    for match in re.finditer(r"标题[^。；;\n]{0,20}逐级递减", text):
        derived_rules.append(
            {
                "type": "heading_font_size_sequence",
                "target": "headings",
                "operation": "decrease_by_level",
                "source": source.path,
                "source_text": context_snippet(text, match.start(), match.end()),
            }
        )

    return derived_rules


def choose_candidates(spec: dict[str, Any], candidates: dict[str, list[FieldCandidate]]) -> list[dict[str, Any]]:
    conflicts = []
    for field, field_candidates in sorted(candidates.items()):
        unique: dict[str, FieldCandidate] = {}
        for candidate in sorted(field_candidates, key=lambda item: item.confidence, reverse=True):
            key = json.dumps(candidate.value, ensure_ascii=False, sort_keys=True)
            unique.setdefault(key, candidate)

        if not unique:
            continue
        if len(unique) == 1:
            selected = next(iter(unique.values()))
            set_nested(spec, field, selected.value)
            continue

        values = list(unique.values())
        conflicts.append(
            {
                "field": field,
                "values": [
                    {
                        "value": candidate.value,
                        "source": candidate.source,
                        "confidence": candidate.confidence,
                        "source_text": candidate.source_text,
                    }
                    for candidate in values
                ],
                "resolution": UNKNOWN,
                "needs_clarification": True,
            }
        )
    return conflicts


def collect_unknowns(spec: dict[str, Any]) -> list[dict[str, Any]]:
    required_fields = {
        "metadata.name": "请确认格式规范名称。",
        "metadata.institution": "请确认学校、期刊或单位名称。",
        "metadata.document_type": "请确认文档类型，例如本科毕业论文、硕士论文、期刊论文或公文。",
        "metadata.year": "请确认格式年份或版本。",
        "page.paper_size": "未找到纸张大小要求，请确认是否为 A4。",
        "page.margins_cm.top": "未找到上页边距要求。",
        "page.margins_cm.bottom": "未找到下页边距要求。",
        "page.margins_cm.left": "未找到左页边距要求。",
        "page.margins_cm.right": "未找到右页边距要求。",
        "body.font.east_asia": "未找到正文中文字体要求。",
        "body.font.size_pt": "未找到正文字号要求。",
        "body.line_spacing": "未找到正文行距要求。",
        "body.first_line_indent_chars": "未找到正文首行缩进要求。",
        "body.alignment": "未找到正文对齐方式要求。",
    }
    unknowns = []
    for field, question in required_fields.items():
        if get_nested(spec, field) == UNKNOWN:
            unknowns.append({"field": field, "question": question})
    return unknowns


def build_validation_rules(spec: dict[str, Any]) -> list[dict[str, Any]]:
    rules = []
    body_checks = {}
    for field in ("font.east_asia", "font.latin", "font.size_pt", "alignment", "line_spacing", "first_line_indent_chars"):
        value = get_nested(spec["body"], field)
        if value != UNKNOWN:
            body_checks[field] = value
    if body_checks:
        rules.append({"target": "body.paragraph", "checks": body_checks})

    page_checks = {}
    if spec["page"]["paper_size"] != UNKNOWN:
        page_checks["paper_size"] = spec["page"]["paper_size"]
    margins = {key: value for key, value in spec["page"]["margins_cm"].items() if value != UNKNOWN}
    if margins:
        page_checks["margins_cm"] = margins
    if page_checks:
        rules.append({"target": "document.page", "checks": page_checks})

    for heading in spec["headings"]:
        checks = {}
        for key, value in heading.get("font", {}).items():
            if value != UNKNOWN:
                checks[f"font.{key}"] = value
        if heading.get("alignment") != UNKNOWN:
            checks["alignment"] = heading["alignment"]
        if checks:
            rules.append({"target": f"heading.level_{heading['level']}", "checks": checks})
    return rules


def parse_format_requirements(
    source_paths: Iterable[Path],
    description: str = "",
    name: str = "",
) -> tuple[dict[str, Any], str, dict[str, Any]]:
    sources = [extract_source(path.expanduser().resolve()) for path in source_paths]
    if not sources:
        raise FormatParserError("至少需要提供一个格式要求文件")

    spec = empty_spec()
    spec["metadata"]["description"] = description
    spec["metadata"]["name"] = name or description.strip() or Path(sources[0].path).stem
    spec["metadata"]["sources"] = [
        {
            "path": source.path,
            "type": source.type,
            "paragraph_count": len(source.paragraphs),
            "table_count": len(source.tables),
            "warnings": source.warnings,
        }
        for source in sources
    ]

    candidates: dict[str, list[FieldCandidate]] = {}
    all_derived_rules = []
    for source in sources:
        collect_metadata_candidates(source, description, candidates)
        collect_page_candidates(source, candidates)
        collect_body_candidates(source, candidates)
        all_derived_rules.extend(collect_special_candidates(source, candidates))

    spec["conflicts"] = choose_candidates(spec, candidates)

    headings = []
    seen_levels = set()
    for source in sources:
        for heading in collect_heading_candidates(source):
            if heading["level"] not in seen_levels:
                headings.append(heading)
                seen_levels.add(heading["level"])
    spec["headings"] = headings
    spec["derived_rules"] = all_derived_rules
    spec["unknowns"] = collect_unknowns(spec)
    spec["validation_rules"] = build_validation_rules(spec)

    spec_md = render_spec_markdown(spec)
    report = build_parse_report(sources, spec)
    return spec, spec_md, report


def format_value(value: Any) -> str:
    if isinstance(value, dict):
        return json.dumps(value, ensure_ascii=False)
    if value is True:
        return "是"
    if value is False:
        return "否"
    return str(value)


def render_spec_markdown(spec: dict[str, Any]) -> str:
    metadata = spec["metadata"]
    lines = [
        f"# {metadata['name']}",
        "",
        "## 元数据",
        "",
        f"- 名称：{metadata['name']}",
        f"- 机构：{metadata['institution']}",
        f"- 文档类型：{metadata['document_type']}",
        f"- 年份/版本：{metadata['year']}",
        "",
        "## 来源文件",
        "",
    ]
    for source in metadata["sources"]:
        lines.append(f"- `{source['path']}`（{source['type']}，段落 {source['paragraph_count']}，表格 {source['table_count']}）")

    lines.extend(
        [
            "",
            "## 页面设置",
            "",
            f"- 纸张：{spec['page']['paper_size']}",
            f"- 方向：{spec['page']['orientation']}",
            f"- 页边距：上 {spec['page']['margins_cm']['top']} cm，下 {spec['page']['margins_cm']['bottom']} cm，左 {spec['page']['margins_cm']['left']} cm，右 {spec['page']['margins_cm']['right']} cm",
            f"- 页眉距离：{spec['page']['header_distance_cm']} cm",
            f"- 页脚距离：{spec['page']['footer_distance_cm']} cm",
            "",
            "## 正文",
            "",
            f"- 中文字体：{spec['body']['font']['east_asia']}",
            f"- 西文字体：{spec['body']['font']['latin']}",
            f"- 字号：{spec['body']['font']['size_cn']} / {spec['body']['font']['size_pt']} pt",
            f"- 行距：{spec['body']['line_spacing']}",
            f"- 首行缩进：{spec['body']['first_line_indent_chars']} 字符",
            f"- 对齐：{spec['body']['alignment']}",
            "",
            "## 标题",
            "",
        ]
    )
    if spec["headings"]:
        for heading in spec["headings"]:
            lines.append(
                f"- {heading['level']} 级标题：{heading['font']['east_asia']}，{heading['font']['size_cn']} / {heading['font']['size_pt']} pt，{heading['alignment']}"
            )
    else:
        lines.append("- unknown")

    lines.extend(
        [
            "",
            "## 表格",
            "",
            f"- 样式：{spec['tables']['style']}",
            f"- 表题位置：{spec['tables']['caption_position']}",
            f"- 表题对齐：{spec['tables']['caption_alignment']}",
            "",
            "## 图",
            "",
            f"- 图题位置：{spec['figures']['caption_position']}",
            f"- 图题对齐：{spec['figures']['caption_alignment']}",
            "",
            "## 公式",
            "",
            f"- 字体：{spec['equations']['font']}",
            f"- 对齐：{spec['equations']['alignment']}",
            f"- 编号：{spec['equations']['numbering']}",
            "",
            "## 参考文献",
            "",
            f"- 样式：{spec['references']['style']}",
            f"- 对齐：{spec['references']['alignment']}",
            f"- 缩进：{spec['references']['indent']}",
            "",
            "## 推导式规则",
            "",
        ]
    )
    if spec["derived_rules"]:
        for rule in spec["derived_rules"]:
            lines.append(f"- {rule['type']}：{rule.get('source_text', '')}")
    else:
        lines.append("- 无")

    lines.extend(["", "## 冲突", ""])
    if spec["conflicts"]:
        for conflict in spec["conflicts"]:
            values = ", ".join(format_value(item["value"]) for item in conflict["values"])
            lines.append(f"- {conflict['field']}：{values}")
    else:
        lines.append("- 无")

    lines.extend(["", "## 待澄清", ""])
    if spec["unknowns"]:
        for unknown in spec["unknowns"]:
            lines.append(f"- {unknown['field']}：{unknown['question']}")
    else:
        lines.append("- 无")
    lines.append("")
    return "\n".join(lines)


def build_parse_report(sources: list[ExtractedSource], spec: dict[str, Any]) -> dict[str, Any]:
    warnings = []
    for source in sources:
        for warning in source.warnings:
            warnings.append({"source": source.path, "warning": warning})
    return {
        "sources": copy.deepcopy(spec["metadata"]["sources"]),
        "extractors": sorted({source.type for source in sources}),
        "warnings": warnings,
        "conflicts": copy.deepcopy(spec["conflicts"]),
        "unknowns": copy.deepcopy(spec["unknowns"]),
    }


def write_outputs(output_dir: Path, spec: dict[str, Any], spec_md: str, report: dict[str, Any]) -> None:
    output_dir.mkdir(parents=True, exist_ok=True)
    (output_dir / "format_spec.md").write_text(spec_md, encoding="utf-8")
    (output_dir / "format_spec.json").write_text(json.dumps(spec, ensure_ascii=False, indent=2), encoding="utf-8")
    (output_dir / "parse_report.json").write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Parse format requirement files into format_spec.md/json.")
    subparsers = parser.add_subparsers(dest="command", required=True)

    parse = subparsers.add_parser("parse", help="parse requirement files")
    parse.add_argument("--output-dir", type=Path, required=True)
    parse.add_argument("--description", default="")
    parse.add_argument("--name", default="")
    parse.add_argument("sources", nargs="+", type=Path)
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    if args.command == "parse":
        spec, spec_md, report = parse_format_requirements(
            source_paths=args.sources,
            description=args.description,
            name=args.name,
        )
        write_outputs(args.output_dir, spec, spec_md, report)
        print(f"已生成: {args.output_dir / 'format_spec.md'}")
        print(f"已生成: {args.output_dir / 'format_spec.json'}")
        print(f"已生成: {args.output_dir / 'parse_report.json'}")
        return
    raise FormatParserError(f"未知命令: {args.command}")


if __name__ == "__main__":
    main()
