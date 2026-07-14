#!/usr/bin/env python3
"""Generic Word formatting engine for format packages."""

from __future__ import annotations

import argparse
import copy
import json
import re
import shutil
import tempfile
import traceback
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from openxml_patches import apply_requested_patches


UNKNOWN = "unknown"

ALIGNMENTS = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    "both": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

TABLE_VERTICAL_ALIGNMENTS = {
    "top": WD_ALIGN_VERTICAL.TOP,
    "center": WD_ALIGN_VERTICAL.CENTER,
    "bottom": WD_ALIGN_VERTICAL.BOTTOM,
}

PAGE_SIZES_CM = {
    "A4": (21.0, 29.7),
    "LETTER": (21.59, 27.94),
}


def is_known(value: Any) -> bool:
    return value not in (None, "", UNKNOWN)


def load_spec(path: Path) -> dict[str, Any]:
    with path.open("r", encoding="utf-8") as handle:
        spec = json.load(handle)
    if not isinstance(spec, dict):
        raise ValueError("format_spec.json 顶层必须是 JSON 对象")
    return spec


def write_report(path: Path | None, report: dict[str, Any]) -> None:
    if path is None:
        return
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")


def empty_report(input_path: Path, output_path: Path, spec_path: Path | None) -> dict[str, Any]:
    return {
        "status": "success",
        "input": str(input_path),
        "output": str(output_path),
        "spec": str(spec_path) if spec_path else None,
        "applied": [],
        "skipped": [],
        "skipped_patches": [],
        "errors": [],
    }


def set_rfont(run, key: str, value: str) -> None:
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn(key), value)


def set_run_font(run, fmt: dict[str, Any]) -> None:
    font = run.font
    if is_known(fmt.get("font_name")):
        font.name = str(fmt["font_name"])
        set_rfont(run, "w:ascii", str(fmt["font_name"]))
        set_rfont(run, "w:hAnsi", str(fmt["font_name"]))

    east_asia = fmt.get("east_asia_font_name") or fmt.get("font_name")
    if is_known(east_asia):
        set_rfont(run, "w:eastAsia", str(east_asia))

    if is_known(fmt.get("font_size_pt")):
        font.size = Pt(float(fmt["font_size_pt"]))
    if is_known(fmt.get("bold")):
        font.bold = bool(fmt["bold"])
    if is_known(fmt.get("italic")):
        font.italic = bool(fmt["italic"])
    if is_known(fmt.get("underline")):
        font.underline = bool(fmt["underline"])


def set_paragraph_format(paragraph, fmt: dict[str, Any], report: dict[str, Any], context: str) -> None:
    paragraph_format = paragraph.paragraph_format

    if is_known(fmt.get("alignment")):
        alignment = ALIGNMENTS.get(str(fmt["alignment"]).lower())
        if alignment is None:
            report["skipped"].append(f"{context}: unsupported alignment {fmt['alignment']}")
        else:
            paragraph.alignment = alignment

    if is_known(fmt.get("line_spacing")):
        paragraph_format.line_spacing = float(fmt["line_spacing"])
    if is_known(fmt.get("first_line_indent_cm")):
        paragraph_format.first_line_indent = Cm(float(fmt["first_line_indent_cm"]))
    if is_known(fmt.get("first_line_indent_chars")):
        # Approximate 2 Chinese chars as 0.74cm, matching the existing v0.1 example.
        paragraph_format.first_line_indent = Cm(float(fmt["first_line_indent_chars"]) * 0.37)
    if is_known(fmt.get("left_indent_cm")):
        paragraph_format.left_indent = Cm(float(fmt["left_indent_cm"]))
    if is_known(fmt.get("right_indent_cm")):
        paragraph_format.right_indent = Cm(float(fmt["right_indent_cm"]))
    if is_known(fmt.get("space_before_pt")):
        paragraph_format.space_before = Pt(float(fmt["space_before_pt"]))
    if is_known(fmt.get("space_after_pt")):
        paragraph_format.space_after = Pt(float(fmt["space_after_pt"]))

    if not paragraph.runs:
        paragraph.add_run()
    for run in paragraph.runs:
        set_run_font(run, fmt)


def legacy_default_format(spec: dict[str, Any]) -> dict[str, Any]:
    return copy.deepcopy(spec.get("default", {}))


def module1_body_format(spec: dict[str, Any]) -> dict[str, Any]:
    body = spec.get("body", {})
    font = body.get("font", {}) if isinstance(body.get("font"), dict) else {}
    return {
        "font_name": font.get("latin"),
        "east_asia_font_name": font.get("east_asia"),
        "font_size_pt": font.get("size_pt"),
        "alignment": body.get("alignment"),
        "line_spacing": body.get("line_spacing"),
        "first_line_indent_chars": body.get("first_line_indent_chars"),
    }


def body_format(spec: dict[str, Any]) -> dict[str, Any]:
    if "default" in spec:
        return legacy_default_format(spec)
    return module1_body_format(spec)


def merge_format(base: dict[str, Any], override: dict[str, Any]) -> dict[str, Any]:
    merged = copy.deepcopy(base)
    for key, value in override.items():
        if is_known(value):
            merged[key] = value
    return merged


def legacy_rule_matches(paragraph, index: int, rule: dict[str, Any]) -> bool:
    match = rule.get("match", {})
    text = paragraph.text.strip()

    if "paragraph_index" in match and int(match["paragraph_index"]) != index:
        return False
    if match.get("non_empty") is True and not text:
        return False
    if "starts_with" in match:
        prefixes = match["starts_with"]
        if isinstance(prefixes, str):
            prefixes = [prefixes]
        if not any(text.startswith(prefix) for prefix in prefixes):
            return False
    if "contains" in match and str(match["contains"]) not in text:
        return False
    if "regex" in match and not re.search(str(match["regex"]), text):
        return False
    return True


def legacy_paragraph_format(spec: dict[str, Any], paragraph, index: int) -> dict[str, Any] | None:
    fmt = body_format(spec)
    for rule in spec.get("rules", []):
        if legacy_rule_matches(paragraph, index, rule):
            return merge_format(fmt, rule.get("format", {}))
    return fmt


def infer_heading_level(text: str) -> int | None:
    stripped = text.strip()
    if not stripped:
        return None
    if re.match(r"^(第[一二三四五六七八九十百]+[章节]|[一二三四五六七八九十]+、|\d+\s+)", stripped):
        return 1
    if re.match(r"^(\d+\.\d+|（[一二三四五六七八九十]+）)", stripped):
        return 2
    if re.match(r"^(\d+\.\d+\.\d+|[（(]\d+[）)])", stripped):
        return 3
    return None


def heading_format_from_rule(rule: dict[str, Any]) -> dict[str, Any]:
    font = rule.get("font", {}) if isinstance(rule.get("font"), dict) else {}
    return {
        "font_name": font.get("latin"),
        "east_asia_font_name": font.get("east_asia"),
        "font_size_pt": font.get("size_pt"),
        "alignment": rule.get("alignment"),
        "bold": rule.get("bold"),
        "first_line_indent_cm": 0,
    }


def module1_paragraph_format(spec: dict[str, Any], paragraph) -> tuple[dict[str, Any], str]:
    level = infer_heading_level(paragraph.text)
    if level is not None:
        for rule in spec.get("headings", []):
            if int(rule.get("level", 0)) == level:
                return merge_format(body_format(spec), heading_format_from_rule(rule)), f"heading.level_{level}"
    return body_format(spec), "body"


def paragraph_is_in_table(paragraph) -> bool:
    parent = paragraph._element.getparent()
    while parent is not None:
        if parent.tag.endswith("}tc"):
            return True
        parent = parent.getparent()
    return False


def apply_page_settings(document, spec: dict[str, Any], report: dict[str, Any]) -> None:
    page = spec.get("page", {})
    if not isinstance(page, dict) or not page:
        report["skipped"].append("page: missing page spec")
        return

    size_name = page.get("size", page.get("paper_size"))
    orientation = page.get("orientation")
    margins = page.get("margins_cm", {})
    applied = False

    for section in document.sections:
        if is_known(orientation):
            orientation_value = str(orientation).lower()
            if orientation_value == "landscape":
                section.orientation = WD_ORIENT.LANDSCAPE
                applied = True
            elif orientation_value == "portrait":
                section.orientation = WD_ORIENT.PORTRAIT
                applied = True
            else:
                report["skipped"].append(f"page.orientation: unsupported {orientation}")
        if is_known(size_name) and str(size_name).upper() in PAGE_SIZES_CM:
            width_cm, height_cm = PAGE_SIZES_CM[str(size_name).upper()]
            if str(orientation).lower() == "landscape":
                width_cm, height_cm = height_cm, width_cm
            section.page_width = Cm(width_cm)
            section.page_height = Cm(height_cm)
            applied = True
        elif is_known(size_name):
            report["skipped"].append(f"page.size: unsupported {size_name}")

        if isinstance(margins, dict):
            for side, attr in (
                ("top", "top_margin"),
                ("bottom", "bottom_margin"),
                ("left", "left_margin"),
                ("right", "right_margin"),
            ):
                value = margins.get(side)
                if is_known(value):
                    setattr(section, attr, Cm(float(value)))
                    applied = True
        if is_known(page.get("header_distance_cm")):
            section.header_distance = Cm(float(page["header_distance_cm"]))
            applied = True
        if is_known(page.get("footer_distance_cm")):
            section.footer_distance = Cm(float(page["footer_distance_cm"]))
            applied = True

    if applied:
        report["applied"].append("page")
    else:
        report["skipped"].append("page: no known values")


def apply_paragraph_settings(document, spec: dict[str, Any], report: dict[str, Any]) -> None:
    body_count = 0
    heading_count = 0
    for index, paragraph in enumerate(document.paragraphs):
        if paragraph_is_in_table(paragraph):
            continue
        if "default" in spec:
            fmt = legacy_paragraph_format(spec, paragraph, index)
            context = "legacy.paragraph"
            if fmt is None:
                continue
        else:
            fmt, context = module1_paragraph_format(spec, paragraph)
        set_paragraph_format(paragraph, fmt, report, context)
        if context.startswith("heading"):
            heading_count += 1
        else:
            body_count += 1

    if body_count:
        report["applied"].append(f"body:{body_count}")
    if heading_count:
        report["applied"].append(f"headings:{heading_count}")


def table_format(spec: dict[str, Any]) -> dict[str, Any]:
    tables = spec.get("tables", {})
    if "default" in spec:
        return {
            "font_name": tables.get("font_name", spec.get("default", {}).get("font_name")),
            "east_asia_font_name": tables.get("east_asia_font_name", spec.get("default", {}).get("east_asia_font_name")),
            "font_size_pt": tables.get("font_size_pt", spec.get("default", {}).get("font_size_pt")),
            "alignment": tables.get("alignment", "center"),
            "line_spacing": tables.get("line_spacing", 1.0),
            "first_line_indent_cm": tables.get("first_line_indent_cm", 0),
            "space_before_pt": tables.get("space_before_pt", 0),
            "space_after_pt": tables.get("space_after_pt", 0),
            "bold_header_row": tables.get("bold_header_row", False),
            "cell_vertical_alignment": tables.get("cell_vertical_alignment", "center"),
        }

    font = tables.get("font", {}) if isinstance(tables.get("font"), dict) else {}
    return {
        "font_name": font.get("latin"),
        "east_asia_font_name": font.get("east_asia"),
        "font_size_pt": font.get("size_pt"),
        "alignment": tables.get("caption_alignment", "center"),
        "line_spacing": 1.0,
        "first_line_indent_cm": 0,
        "bold_header_row": tables.get("bold_header_row", True),
        "cell_vertical_alignment": "center",
    }


def apply_table_settings(document, spec: dict[str, Any], report: dict[str, Any]) -> None:
    if not document.tables:
        report["skipped"].append("tables: no tables")
        return
    fmt = table_format(spec)
    vertical_alignment = TABLE_VERTICAL_ALIGNMENTS.get(str(fmt.get("cell_vertical_alignment", "center")).lower())
    cell_count = 0
    for table in document.tables:
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                if vertical_alignment is not None:
                    cell.vertical_alignment = vertical_alignment
                cell_fmt = copy.deepcopy(fmt)
                if row_index == 0 and fmt.get("bold_header_row"):
                    cell_fmt["bold"] = True
                for paragraph in cell.paragraphs:
                    set_paragraph_format(paragraph, cell_fmt, report, "table.cell")
                cell_count += 1
    report["applied"].append(f"tables:{cell_count}")


def apply_spec_to_document(document, spec: dict[str, Any], report: dict[str, Any]) -> None:
    apply_page_settings(document, spec, report)
    apply_paragraph_settings(document, spec, report)
    apply_table_settings(document, spec, report)


def format_document(
    input_path: Path,
    output_path: Path,
    spec: dict[str, Any],
    report_path: Path | None = None,
    spec_path: Path | None = None,
) -> dict[str, Any]:
    input_path = input_path.resolve()
    output_path = output_path.resolve()
    report = empty_report(input_path, output_path, spec_path)

    try:
        if input_path.suffix.lower() != ".docx":
            raise ValueError("当前工具只支持 .docx；旧版 .doc 需先转换成 .docx。")
        if not input_path.is_file():
            raise FileNotFoundError(f"输入文件不存在: {input_path}")

        output_path.parent.mkdir(parents=True, exist_ok=True)
        with tempfile.TemporaryDirectory(prefix="word_format_engine_") as tmp:
            tmp_docx = Path(tmp) / "working.docx"
            shutil.copyfile(input_path, tmp_docx)
            document = Document(str(tmp_docx))
            apply_spec_to_document(document, spec, report)
            document.save(str(tmp_docx))

            patch_result = apply_requested_patches(tmp_docx, spec)
            report["skipped_patches"].extend(patch_result["skipped"])
            report["skipped_patches"].extend(patch_result["unknown"])
            if patch_result["applied"]:
                report["applied"].extend(f"patch:{name}" for name in patch_result["applied"])

            shutil.copyfile(tmp_docx, output_path)
    except Exception as exc:
        report["status"] = "failed"
        report["errors"].append({"message": str(exc), "traceback": traceback.format_exc()})
        write_report(report_path, report)
        return report

    write_report(report_path, report)
    return report


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Format a .docx file according to format_spec.json.")
    parser.add_argument("--spec", type=Path, required=True, help="format_spec.json")
    parser.add_argument("--input", type=Path, required=True, help="原始 .docx 文件")
    parser.add_argument("--output", type=Path, required=True, help="输出 .docx 文件")
    parser.add_argument("--report", type=Path, default=None, help="输出格式化报告 JSON")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    spec = load_spec(args.spec)
    report = format_document(args.input, args.output, spec, args.report, args.spec)
    if report["status"] == "failed":
        raise SystemExit(1)
    print(f"已生成: {args.output}")
    if args.report:
        print(f"已生成: {args.report}")


if __name__ == "__main__":
    main()
