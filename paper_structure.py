#!/usr/bin/env python3
"""Analyze logical structure from a raw thesis .docx."""

from __future__ import annotations

import argparse
import json
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from docx import Document


class PaperStructureError(ValueError):
    """Raised when a paper structure input cannot be analyzed."""


@dataclass
class ParagraphInfo:
    index: int
    text: str
    style: str
    has_drawing: bool = False
    has_math: bool = False


@dataclass
class TableInfo:
    index: int
    rows: int
    columns: int
    preview: str


@dataclass
class AnalyzeState:
    in_references: bool = False
    reference_start: int | None = None
    reference_count: int = 0
    warnings: list[str] = field(default_factory=list)


def normalize_text(text: str) -> str:
    return re.sub(r"\s+", " ", text.replace("\u3000", " ")).strip()


def paragraph_has_drawing(paragraph) -> bool:
    return paragraph._element.find(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing") is not None


def paragraph_has_math(paragraph) -> bool:
    return paragraph._element.find(".//{http://schemas.openxmlformats.org/officeDocument/2006/math}oMath") is not None


def read_paragraphs(document: Document) -> list[ParagraphInfo]:
    paragraphs = []
    for index, paragraph in enumerate(document.paragraphs):
        style_name = paragraph.style.name if paragraph.style is not None else ""
        paragraphs.append(
            ParagraphInfo(
                index=index,
                text=normalize_text(paragraph.text),
                style=style_name,
                has_drawing=paragraph_has_drawing(paragraph),
                has_math=paragraph_has_math(paragraph),
            )
        )
    return paragraphs


def read_tables(document: Document) -> list[TableInfo]:
    tables = []
    for index, table in enumerate(document.tables):
        rows = len(table.rows)
        columns = len(table.columns)
        preview_cells = []
        for row in table.rows[:2]:
            for cell in row.cells[:3]:
                text = normalize_text(cell.text)
                if text:
                    preview_cells.append(text)
        tables.append(
            TableInfo(
                index=index,
                rows=rows,
                columns=columns,
                preview=" | ".join(preview_cells[:6]),
            )
        )
    return tables


def style_heading_level(style_name: str) -> int | None:
    style = style_name.strip().lower()
    patterns = [
        (1, [r"heading\s*1", r"标题\s*1", r"标题\s*一"]),
        (2, [r"heading\s*2", r"标题\s*2", r"标题\s*二"]),
        (3, [r"heading\s*3", r"标题\s*3", r"标题\s*三"]),
    ]
    for level, level_patterns in patterns:
        if any(re.search(pattern, style) for pattern in level_patterns):
            return level
    return None


def text_heading_level(text: str) -> int | None:
    if not text:
        return None
    if re.match(r"^(第[一二三四五六七八九十百\d]+[章节篇]|[一二三四五六七八九十]+、|\d+\s+[^.\d])", text):
        return 1
    if re.match(r"^(\d+\.\d+|（[一二三四五六七八九十]+）|\([一二三四五六七八九十]+\))", text):
        return 2
    if re.match(r"^(\d+\.\d+\.\d+|（\d+）|\(\d+\)|\d+[）)])", text):
        return 3
    return None


def infer_heading_level(paragraph: ParagraphInfo) -> tuple[int | None, float, str]:
    style_level = style_heading_level(paragraph.style)
    if style_level is not None:
        return style_level, 0.95, "style"
    text_level = text_heading_level(paragraph.text)
    if text_level is not None:
        return text_level, 0.8, "text_pattern"
    return None, 0.0, ""


def source_for_paragraph(paragraph: ParagraphInfo) -> dict[str, Any]:
    return {
        "kind": "paragraph",
        "index": paragraph.index,
        "style": paragraph.style,
    }


def make_block(
    block_id: str,
    block_type: str,
    paragraph: ParagraphInfo,
    confidence: float,
    **extra: Any,
) -> dict[str, Any]:
    block = {
        "id": block_id,
        "type": block_type,
        "text": paragraph.text,
        "source": source_for_paragraph(paragraph),
        "confidence": confidence,
    }
    block.update(extra)
    return block


def is_exact_heading(text: str, names: set[str]) -> bool:
    stripped = re.sub(r"[\s:：]+", "", text)
    return stripped in names


def classify_special(paragraph: ParagraphInfo) -> tuple[str | None, float]:
    text = paragraph.text
    if not text:
        return None, 0.0
    if is_exact_heading(text, {"摘要", "中文摘要"}):
        return "abstract_heading", 0.9
    if is_exact_heading(text, {"abstract", "英文摘要"}):
        return "abstract_heading", 0.85
    if re.match(r"^关键词[:：]", text):
        return "keywords", 0.9
    if is_exact_heading(text, {"目录"}):
        return "toc_heading", 0.9
    if is_exact_heading(text, {"参考文献", "references"}):
        return "references_heading", 0.95
    if is_exact_heading(text, {"致谢"}):
        return "acknowledgements_heading", 0.9
    if re.match(r"^附录", text):
        return "appendix_heading", 0.85
    return None, 0.0


def parse_caption_number(text: str, prefix: str) -> str | None:
    pattern = rf"^{prefix}\s*([0-9]+(?:[-.．][0-9]+)*|[一二三四五六七八九十]+(?:[-.．][一二三四五六七八九十]+)*)"
    match = re.match(pattern, text, re.IGNORECASE)
    return match.group(1) if match else None


def is_figure_caption(text: str) -> bool:
    return bool(re.match(r"^(图\s*[0-9一二三四五六七八九十]+|figure\s+\d+)", text, re.IGNORECASE))


def is_table_caption(text: str) -> bool:
    return bool(re.match(r"^(表\s*[0-9一二三四五六七八九十]+|table\s+\d+)", text, re.IGNORECASE))


def is_reference_item(text: str) -> bool:
    return bool(re.match(r"^(\[\d+\]|\d+[.．、])\s*", text))


def nearest_table_index_after(paragraph_index: int, paragraph_count: int, tables: list[TableInfo]) -> int | None:
    # python-docx does not expose exact document-order table indices in a simple way.
    # This deterministic approximation is enough for module 3 v1 reporting.
    if not tables:
        return None
    ratio = paragraph_index / max(paragraph_count, 1)
    guessed = min(int(ratio * len(tables)), len(tables) - 1)
    return tables[guessed].index


def analyze_paragraphs(paragraphs: list[ParagraphInfo], tables: list[TableInfo]) -> dict[str, Any]:
    blocks = []
    sections = []
    figures = []
    table_captions = []
    references = []
    preserve = []
    state = AnalyzeState()

    for paragraph in paragraphs:
        block_id = f"p{paragraph.index:04d}"

        if not paragraph.text and not paragraph.has_drawing and not paragraph.has_math:
            continue

        if paragraph.has_drawing:
            block = make_block(block_id, "preserve", paragraph, 0.6, reason="drawing")
            blocks.append(block)
            preserve.append(block)
            continue

        if paragraph.has_math:
            block = make_block(block_id, "preserve", paragraph, 0.6, reason="math")
            blocks.append(block)
            preserve.append(block)
            continue

        special_type, special_confidence = classify_special(paragraph)
        if special_type:
            block = make_block(block_id, special_type, paragraph, special_confidence)
            blocks.append(block)
            if special_type == "references_heading":
                state.in_references = True
                state.reference_start = paragraph.index
            elif special_type in {"acknowledgements_heading", "appendix_heading"}:
                state.in_references = False
            continue

        if state.in_references:
            level, _, _ = infer_heading_level(paragraph)
            if level == 1 and not is_reference_item(paragraph.text):
                state.in_references = False
            else:
                item = make_block(block_id, "reference_item", paragraph, 0.8 if is_reference_item(paragraph.text) else 0.6)
                blocks.append(item)
                references.append(item)
                state.reference_count += 1
                continue

        if is_figure_caption(paragraph.text):
            number = parse_caption_number(paragraph.text, "图") or parse_caption_number(paragraph.text, "figure")
            figure = make_block(block_id, "figure_caption", paragraph, 0.85, number=number)
            blocks.append(figure)
            figures.append(figure)
            continue

        if is_table_caption(paragraph.text):
            number = parse_caption_number(paragraph.text, "表") or parse_caption_number(paragraph.text, "table")
            table_index = nearest_table_index_after(paragraph.index, len(paragraphs), tables)
            caption = make_block(block_id, "table_caption", paragraph, 0.85, number=number, table_index=table_index)
            blocks.append(caption)
            table_captions.append(caption)
            continue

        level, confidence, method = infer_heading_level(paragraph)
        if level is not None:
            heading = make_block(block_id, "heading", paragraph, confidence, level=level, method=method)
            blocks.append(heading)
            sections.append(
                {
                    "id": block_id,
                    "level": level,
                    "title": paragraph.text,
                    "source": source_for_paragraph(paragraph),
                    "confidence": confidence,
                    "method": method,
                }
            )
            continue

        body = make_block(block_id, "body", paragraph, 0.55)
        blocks.append(body)
        preserve_item = make_block(block_id, "preserve", paragraph, 0.4, reason="unclassified_paragraph")
        preserve.append(preserve_item)

    return {
        "blocks": blocks,
        "sections": sections,
        "figures": figures,
        "table_captions": table_captions,
        "references": references,
        "preserve": preserve,
        "reference_start": state.reference_start,
    }


def table_entries(tables: list[TableInfo]) -> list[dict[str, Any]]:
    return [
        {
            "id": f"tbl{table.index:04d}",
            "type": "table",
            "rows": table.rows,
            "columns": table.columns,
            "preview": table.preview,
            "source": {"kind": "table", "index": table.index},
            "confidence": 1.0,
        }
        for table in tables
    ]


def build_structure(input_path: Path, document: Document) -> dict[str, Any]:
    paragraphs = read_paragraphs(document)
    tables = read_tables(document)
    analyzed = analyze_paragraphs(paragraphs, tables)
    table_data = table_entries(tables)
    warnings = []
    if not analyzed["sections"]:
        warnings.append("未识别到章节标题")

    return {
        "metadata": {
            "input": str(input_path),
            "paragraph_count": len(paragraphs),
            "table_count": len(tables),
        },
        "blocks": analyzed["blocks"] + table_data,
        "sections": analyzed["sections"],
        "figures": analyzed["figures"],
        "tables": table_data,
        "table_captions": analyzed["table_captions"],
        "references": analyzed["references"],
        "preserve": analyzed["preserve"],
        "warnings": warnings,
    }


def build_report(structure: dict[str, Any]) -> dict[str, Any]:
    counts = {
        "paragraphs": structure["metadata"]["paragraph_count"],
        "tables": structure["metadata"]["table_count"],
        "headings": len(structure["sections"]),
        "figures": len(structure["figures"]),
        "table_captions": len(structure["table_captions"]),
        "references": len(structure["references"]),
        "preserve": len(structure["preserve"]),
    }
    return {
        "status": "success",
        "input": structure["metadata"]["input"],
        "counts": counts,
        "warnings": structure["warnings"],
    }


def render_markdown(structure: dict[str, Any], report: dict[str, Any]) -> str:
    lines = [
        "# 论文结构识别结果",
        "",
        "## 文档概况",
        "",
        f"- 段落数：{report['counts']['paragraphs']}",
        f"- 表格数：{report['counts']['tables']}",
        f"- 标题数：{report['counts']['headings']}",
        f"- 图题数：{report['counts']['figures']}",
        f"- 表题数：{report['counts']['table_captions']}",
        f"- 参考文献条目：{report['counts']['references']}",
        f"- 保留项：{report['counts']['preserve']}",
        "",
        "## 章节结构",
        "",
    ]

    if structure["sections"]:
        for section in structure["sections"]:
            indent = "  " * max(section["level"] - 1, 0)
            lines.append(f"{indent}- H{section['level']} p{section['source']['index']:04d} {section['title']}")
    else:
        lines.append("- 未识别到章节标题")

    lines.extend(["", "## 图", ""])
    if structure["figures"]:
        for figure in structure["figures"]:
            lines.append(f"- p{figure['source']['index']:04d} {figure['text']}")
    else:
        lines.append("- 无")

    lines.extend(["", "## 表", ""])
    if structure["table_captions"]:
        for caption in structure["table_captions"]:
            table_index = caption.get("table_index")
            suffix = f" -> table {table_index}" if table_index is not None else ""
            lines.append(f"- p{caption['source']['index']:04d} {caption['text']}{suffix}")
    else:
        lines.append("- 无")

    lines.extend(["", "## 参考文献", ""])
    if structure["references"]:
        for reference in structure["references"][:20]:
            lines.append(f"- p{reference['source']['index']:04d} {reference['text']}")
        if len(structure["references"]) > 20:
            lines.append(f"- ... 共 {len(structure['references'])} 条")
    else:
        lines.append("- 无")

    lines.extend(["", "## 保留项", ""])
    if structure["preserve"]:
        for item in structure["preserve"][:30]:
            lines.append(f"- p{item['source']['index']:04d}：{item.get('reason', item['type'])}")
        if len(structure["preserve"]) > 30:
            lines.append(f"- ... 共 {len(structure['preserve'])} 项")
    else:
        lines.append("- 无")

    if structure["warnings"]:
        lines.extend(["", "## 警告", ""])
        for warning in structure["warnings"]:
            lines.append(f"- {warning}")

    lines.append("")
    return "\n".join(lines)


def analyze_docx(input_path: Path) -> tuple[dict[str, Any], str, dict[str, Any]]:
    if input_path.suffix.lower() != ".docx":
        raise PaperStructureError("当前只支持 .docx")
    if not input_path.is_file():
        raise PaperStructureError(f"输入文件不存在: {input_path}")
    document = Document(str(input_path))
    structure = build_structure(input_path.resolve(), document)
    report = build_report(structure)
    markdown = render_markdown(structure, report)
    return structure, markdown, report


def write_outputs(output_json: Path, output_md: Path, report_path: Path, structure: dict[str, Any], markdown: str, report: dict[str, Any]) -> None:
    output_json.parent.mkdir(parents=True, exist_ok=True)
    output_md.parent.mkdir(parents=True, exist_ok=True)
    report_path.parent.mkdir(parents=True, exist_ok=True)
    output_json.write_text(json.dumps(structure, ensure_ascii=False, indent=2), encoding="utf-8")
    output_md.write_text(markdown, encoding="utf-8")
    report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Analyze logical structure from a raw thesis .docx.")
    subparsers = parser.add_subparsers(dest="command", required=True)

    analyze = subparsers.add_parser("analyze", help="analyze a raw .docx")
    analyze.add_argument("--input", type=Path, required=True)
    analyze.add_argument("--output-md", type=Path, required=True)
    analyze.add_argument("--output-json", type=Path, required=True)
    analyze.add_argument("--report", type=Path, required=True)
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    if args.command == "analyze":
        structure, markdown, report = analyze_docx(args.input)
        write_outputs(args.output_json, args.output_md, args.report, structure, markdown, report)
        print(f"已生成: {args.output_md}")
        print(f"已生成: {args.output_json}")
        print(f"已生成: {args.report}")
        return
    raise PaperStructureError(f"未知命令: {args.command}")


if __name__ == "__main__":
    main()
