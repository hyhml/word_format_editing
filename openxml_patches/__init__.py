"""OpenXML patch registry and conservative patch implementations.

The generic engine handles page, paragraph, heading, and table text formatting
with python-docx. This module handles rules that need lower-level Wordprocessing
ML changes or document-wide post-processing.
"""

from __future__ import annotations

import re
import shutil
import tempfile
import zipfile
from pathlib import Path
from typing import Any, Callable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from lxml import etree


UNKNOWN = "unknown"
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"

KNOWN_PATCHES = {
    "header_footer",
    "three_line_table",
    "captions",
    "equation_numbering",
    "math_font",
    "references",
}

ALIGNMENTS = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    "both": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def is_known(value: Any) -> bool:
    return value not in (None, "", UNKNOWN)


def requested_patches(spec: dict[str, Any]) -> list[str]:
    raw = spec.get("openxml_patches", None)
    if raw is True:
        return sorted(KNOWN_PATCHES)
    if raw is False:
        return []
    if raw is None:
        return inferred_patches(spec)
    if not raw:
        return []
    if not isinstance(raw, list):
        return [str(raw)]
    return [str(item) for item in raw]


def inferred_patches(spec: dict[str, Any]) -> list[str]:
    patches = []
    tables = spec.get("tables", {}) if isinstance(spec.get("tables"), dict) else {}
    figures = spec.get("figures", {}) if isinstance(spec.get("figures"), dict) else {}
    equations = spec.get("equations", {}) if isinstance(spec.get("equations"), dict) else {}
    references = spec.get("references", {}) if isinstance(spec.get("references"), dict) else {}
    headers_footers = spec.get("headers_footers", {}) if isinstance(spec.get("headers_footers"), dict) else {}

    if tables.get("style") == "three_line":
        patches.append("three_line_table")
    if any_known(headers_footers):
        patches.append("header_footer")
    if any_known(figures) or is_known(tables.get("caption_position")) or is_known(tables.get("caption_alignment")):
        patches.append("captions")
    if any_known(references):
        patches.append("references")
    if is_known(equations.get("font")):
        patches.append("math_font")
    if is_known(equations.get("numbering")):
        patches.append("equation_numbering")
    return patches


def any_known(value: Any) -> bool:
    if isinstance(value, dict):
        return any(any_known(item) for item in value.values())
    if isinstance(value, list):
        return any(any_known(item) for item in value)
    return is_known(value)


def apply_requested_patches(docx_path: Path, spec: dict[str, Any]) -> dict[str, list[str]]:
    result = {"applied": [], "skipped": [], "unknown": [], "errors": []}
    patch_functions: dict[str, Callable[[Path, dict[str, Any]], str | None]] = {
        "three_line_table": apply_three_line_table,
        "header_footer": apply_header_footer,
        "captions": apply_captions,
        "references": apply_references,
        "math_font": apply_math_font,
        "equation_numbering": apply_equation_numbering,
    }

    for patch_name in requested_patches(spec):
        patch_function = patch_functions.get(patch_name)
        if patch_function is None:
            result["unknown"].append(f"{patch_name}: unknown patch")
            continue
        try:
            message = patch_function(Path(docx_path), spec)
        except Exception as exc:  # keep base formatting output intact
            result["errors"].append(f"{patch_name}: {exc}")
            continue
        if message:
            result["applied"].append(f"{patch_name}: {message}")
        else:
            result["skipped"].append(f"{patch_name}: no applicable content or known rules")
    return result


def ensure_child(parent, tag: str):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def set_border(parent, side: str, val: str, size: str = "0", color: str = "000000") -> None:
    border = ensure_child(parent, f"w:{side}")
    border.set(qn("w:val"), val)
    border.set(qn("w:sz"), size)
    border.set(qn("w:space"), "0")
    border.set(qn("w:color"), color)


def apply_three_line_table(docx_path: Path, spec: dict[str, Any]) -> str | None:
    tables = spec.get("tables", {}) if isinstance(spec.get("tables"), dict) else {}
    explicit = "three_line_table" in requested_patches(spec)
    if tables.get("style") != "three_line" and not explicit:
        return None

    document = Document(str(docx_path))
    if not document.tables:
        return None

    for table in document.tables:
        tbl = table._tbl
        tbl_pr = tbl.tblPr
        if tbl_pr is None:
            tbl_pr = OxmlElement("w:tblPr")
            tbl.insert(0, tbl_pr)
        old_borders = tbl_pr.find(qn("w:tblBorders"))
        if old_borders is not None:
            tbl_pr.remove(old_borders)
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
        set_border(borders, "top", "single", "12")
        set_border(borders, "bottom", "single", "12")
        for side in ("left", "right", "insideH", "insideV"):
            set_border(borders, side, "nil")

        if table.rows:
            for cell in table.rows[0].cells:
                tc_pr = cell._tc.get_or_add_tcPr()
                old_cell_borders = tc_pr.find(qn("w:tcBorders"))
                if old_cell_borders is not None:
                    tc_pr.remove(old_cell_borders)
                cell_borders = OxmlElement("w:tcBorders")
                tc_pr.append(cell_borders)
                set_border(cell_borders, "bottom", "single", "6")

    document.save(str(docx_path))
    return f"{len(document.tables)} tables"


def font_from_spec(spec: dict[str, Any], section_name: str, fallback_to_body: bool = True) -> dict[str, Any]:
    section = spec.get(section_name, {}) if isinstance(spec.get(section_name), dict) else {}
    font = section.get("font", {}) if isinstance(section.get("font"), dict) else {}
    body = spec.get("body", {}) if isinstance(spec.get("body"), dict) else {}
    body_font = body.get("font", {}) if isinstance(body.get("font"), dict) else {}
    legacy_default = spec.get("default", {}) if isinstance(spec.get("default"), dict) else {}
    return {
        "latin": font.get("latin") or (body_font.get("latin") if fallback_to_body else None) or legacy_default.get("font_name"),
        "east_asia": font.get("east_asia") or (body_font.get("east_asia") if fallback_to_body else None) or legacy_default.get("east_asia_font_name"),
        "size_pt": font.get("size_pt") or (body_font.get("size_pt") if fallback_to_body else None) or legacy_default.get("font_size_pt"),
    }


def apply_run_font(run, font: dict[str, Any], bold: Any = None) -> None:
    if is_known(font.get("latin")):
        run.font.name = str(font["latin"])
        r_pr = run._element.get_or_add_rPr()
        r_fonts = r_pr.rFonts
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.append(r_fonts)
        r_fonts.set(qn("w:ascii"), str(font["latin"]))
        r_fonts.set(qn("w:hAnsi"), str(font["latin"]))
    if is_known(font.get("east_asia")):
        r_pr = run._element.get_or_add_rPr()
        r_fonts = r_pr.rFonts
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.append(r_fonts)
        r_fonts.set(qn("w:eastAsia"), str(font["east_asia"]))
    if is_known(font.get("size_pt")):
        run.font.size = Pt(float(font["size_pt"]))
    if is_known(bold):
        run.font.bold = bool(bold)


def apply_paragraph_style(paragraph, alignment: Any, font: dict[str, Any], keep_next: bool = False, first_line_cm: float | None = 0) -> None:
    if is_known(alignment):
        paragraph.alignment = ALIGNMENTS.get(str(alignment).lower(), paragraph.alignment)
    if keep_next:
        paragraph.paragraph_format.keep_with_next = True
        paragraph.paragraph_format.keep_together = True
    if first_line_cm is not None:
        paragraph.paragraph_format.first_line_indent = Cm(first_line_cm)
    if not paragraph.runs:
        paragraph.add_run()
    for run in paragraph.runs:
        apply_run_font(run, font)


def caption_font(spec: dict[str, Any], kind: str) -> dict[str, Any]:
    font = font_from_spec(spec, kind)
    if not is_known(font.get("size_pt")):
        font["size_pt"] = 10.5
    if not is_known(font.get("east_asia")):
        font["east_asia"] = "宋体"
    if not is_known(font.get("latin")):
        font["latin"] = "Times New Roman"
    return font


def is_table_caption(text: str) -> bool:
    return bool(re.match(r"^\s*(表\s*[0-9一二三四五六七八九十]+|table\s+\d+)", text, re.IGNORECASE))


def is_figure_caption(text: str) -> bool:
    return bool(re.match(r"^\s*(图\s*[0-9一二三四五六七八九十]+|figure\s+\d+)", text, re.IGNORECASE))


def apply_captions(docx_path: Path, spec: dict[str, Any]) -> str | None:
    document = Document(str(docx_path))
    tables = spec.get("tables", {}) if isinstance(spec.get("tables"), dict) else {}
    figures = spec.get("figures", {}) if isinstance(spec.get("figures"), dict) else {}
    table_count = 0
    figure_count = 0

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if is_table_caption(text):
            apply_paragraph_style(
                paragraph,
                tables.get("caption_alignment", "center"),
                caption_font(spec, "tables"),
                keep_next=True,
                first_line_cm=0,
            )
            table_count += 1
        elif is_figure_caption(text):
            apply_paragraph_style(
                paragraph,
                figures.get("caption_alignment", "center"),
                caption_font(spec, "figures"),
                keep_next=True,
                first_line_cm=0,
            )
            figure_count += 1

    if not table_count and not figure_count:
        return None
    document.save(str(docx_path))
    return f"{table_count} table captions, {figure_count} figure captions"


def header_footer_has_rules(spec: dict[str, Any]) -> bool:
    headers_footers = spec.get("headers_footers", {}) if isinstance(spec.get("headers_footers"), dict) else {}
    return any_known(headers_footers)


def clear_paragraph(paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def header_font(spec: dict[str, Any]) -> dict[str, Any]:
    headers_footers = spec.get("headers_footers", {}) if isinstance(spec.get("headers_footers"), dict) else {}
    header = headers_footers.get("header", {}) if isinstance(headers_footers.get("header"), dict) else {}
    return {
        "latin": header.get("font") or "Times New Roman",
        "east_asia": header.get("font") or "宋体",
        "size_pt": header.get("size_pt"),
    }


def footer_font(spec: dict[str, Any]) -> dict[str, Any]:
    headers_footers = spec.get("headers_footers", {}) if isinstance(spec.get("headers_footers"), dict) else {}
    footer = headers_footers.get("footer", {}) if isinstance(headers_footers.get("footer"), dict) else {}
    return {
        "latin": footer.get("font") or "Times New Roman",
        "east_asia": footer.get("font") or "宋体",
        "size_pt": footer.get("size_pt") or 9,
    }


def add_page_number_field(paragraph, font: dict[str, Any]) -> None:
    begin_run = paragraph.add_run()
    apply_run_font(begin_run, font)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin_run._r.append(begin)

    instr_run = paragraph.add_run()
    apply_run_font(instr_run, font)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    instr_run._r.append(instr)

    end_run = paragraph.add_run()
    apply_run_font(end_run, font)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def apply_header_footer(docx_path: Path, spec: dict[str, Any]) -> str | None:
    if not header_footer_has_rules(spec):
        return None

    headers_footers = spec.get("headers_footers", {})
    header = headers_footers.get("header", {}) if isinstance(headers_footers.get("header"), dict) else {}
    footer = headers_footers.get("footer", {}) if isinstance(headers_footers.get("footer"), dict) else {}
    header_text = header.get("text")
    footer_page_number = footer.get("page_number")
    applied = 0

    document = Document(str(docx_path))
    for section in document.sections:
        if is_known(header_text):
            section.header.is_linked_to_previous = False
            paragraph = section.header.paragraphs[0]
            clear_paragraph(paragraph)
            paragraph.alignment = ALIGNMENTS.get(str(header.get("alignment", "center")).lower(), WD_ALIGN_PARAGRAPH.CENTER)
            run = paragraph.add_run(str(header_text))
            apply_run_font(run, header_font(spec))
            applied += 1

        if footer_page_number is True or str(footer_page_number).lower() in {"true", "page_number", "page"}:
            section.footer.is_linked_to_previous = False
            paragraph = section.footer.paragraphs[0]
            clear_paragraph(paragraph)
            paragraph.alignment = ALIGNMENTS.get(str(footer.get("alignment", "center")).lower(), WD_ALIGN_PARAGRAPH.CENTER)
            add_page_number_field(paragraph, footer_font(spec))
            applied += 1

    if not applied:
        return None
    document.save(str(docx_path))
    return f"{applied} header/footer parts"


def is_reference_heading(text: str) -> bool:
    stripped = re.sub(r"[\s:：]+", "", text)
    return stripped in {"参考文献", "references"}


def is_stop_heading(text: str) -> bool:
    stripped = re.sub(r"[\s:：]+", "", text)
    return stripped in {"致谢", "附录", "acknowledgements", "appendix"}


def apply_references(docx_path: Path, spec: dict[str, Any]) -> str | None:
    references = spec.get("references", {}) if isinstance(spec.get("references"), dict) else {}
    if not any_known(references) and "references" not in requested_patches(spec):
        return None

    document = Document(str(docx_path))
    in_references = False
    count = 0
    font = font_from_spec(spec, "references")
    alignment = references.get("alignment", "justify")
    hanging = references.get("indent") == "hanging"

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if is_reference_heading(text):
            in_references = True
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
        if in_references and is_stop_heading(text):
            break
        if not in_references or not text:
            continue
        paragraph.alignment = ALIGNMENTS.get(str(alignment).lower(), WD_ALIGN_PARAGRAPH.JUSTIFY)
        if hanging:
            paragraph.paragraph_format.left_indent = Cm(0.74)
            paragraph.paragraph_format.first_line_indent = Cm(-0.74)
        else:
            paragraph.paragraph_format.first_line_indent = Cm(0)
        for run in paragraph.runs:
            apply_run_font(run, font)
            if run.text:
                run.text = run.text.replace("\xa0", " ")
        count += 1

    if not count:
        return None
    document.save(str(docx_path))
    return f"{count} reference paragraphs"


def read_zip_xml(zip_file: zipfile.ZipFile, name: str):
    try:
        return etree.fromstring(zip_file.read(name))
    except KeyError:
        return None


def write_docx_with_xml_updates(docx_path: Path, updates: dict[str, bytes]) -> None:
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
        tmp_path = Path(tmp_file.name)
    try:
        with zipfile.ZipFile(docx_path, "r") as src, zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as dst:
            for item in src.infolist():
                data = updates.get(item.filename)
                if data is None:
                    data = src.read(item.filename)
                dst.writestr(item, data)
        shutil.copyfile(tmp_path, docx_path)
    finally:
        tmp_path.unlink(missing_ok=True)


def apply_math_font(docx_path: Path, spec: dict[str, Any]) -> str | None:
    equations = spec.get("equations", {}) if isinstance(spec.get("equations"), dict) else {}
    font_name = equations.get("font")
    if not is_known(font_name):
        font_name = "Times New Roman" if "math_font" in requested_patches(spec) else None
    if not is_known(font_name):
        return None

    updates: dict[str, bytes] = {}
    math_runs = 0
    with zipfile.ZipFile(docx_path, "r") as package:
        document_root = read_zip_xml(package, "word/document.xml")
        if document_root is not None:
            for math_run in document_root.iter(f"{{{M_NS}}}r"):
                math_runs += 1
                math_r_pr = math_run.find(f"{{{M_NS}}}rPr")
                if math_r_pr is None:
                    math_r_pr = etree.Element(f"{{{M_NS}}}rPr")
                    math_run.insert(0, math_r_pr)
                r_fonts = math_r_pr.find(f"{{{W_NS}}}rFonts")
                if r_fonts is None:
                    r_fonts = etree.SubElement(math_r_pr, f"{{{W_NS}}}rFonts")
                r_fonts.set(f"{{{W_NS}}}ascii", str(font_name))
                r_fonts.set(f"{{{W_NS}}}hAnsi", str(font_name))
                r_fonts.set(f"{{{W_NS}}}cs", str(font_name))
            if math_runs:
                updates["word/document.xml"] = etree.tostring(document_root, encoding="UTF-8", xml_declaration=True)

        settings_root = read_zip_xml(package, "word/settings.xml")
        if settings_root is not None:
            math_pr = settings_root.find(f"{{{M_NS}}}mathPr")
            if math_pr is None:
                math_pr = etree.SubElement(settings_root, f"{{{M_NS}}}mathPr")
            math_font = math_pr.find(f"{{{M_NS}}}mathFont")
            if math_font is None:
                math_font = etree.SubElement(math_pr, f"{{{M_NS}}}mathFont")
            math_font.set(f"{{{M_NS}}}val", str(font_name))
            updates["word/settings.xml"] = etree.tostring(settings_root, encoding="UTF-8", xml_declaration=True)

    if not updates:
        return None
    write_docx_with_xml_updates(docx_path, updates)
    return f"{font_name}, {math_runs} math runs"


def paragraph_has_math(paragraph) -> bool:
    return paragraph._element.find(f".//{{{M_NS}}}oMath") is not None or paragraph._element.find(f".//{{{M_NS}}}oMathPara") is not None


def apply_equation_numbering(docx_path: Path, spec: dict[str, Any]) -> str | None:
    equations = spec.get("equations", {}) if isinstance(spec.get("equations"), dict) else {}
    numbering = equations.get("numbering")
    if not is_known(numbering) and "equation_numbering" not in requested_patches(spec):
        return None

    document = Document(str(docx_path))
    count = 0
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if paragraph_has_math(paragraph) or re.search(r"[（(]\s*\d+(?:[-.]\d+)*\s*[）)]$", text):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            count += 1
    if not count:
        return None
    document.save(str(docx_path))
    return f"{count} equation paragraphs aligned"
