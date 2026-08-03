"""Extract requirement files into source blocks with stable evidence IDs."""

from __future__ import annotations

from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Any


SUPPORTED_SUFFIXES = {".txt", ".md", ".json", ".docx", ".pdf"}


@dataclass
class SourceBlock:
    id: str
    source_id: str
    type: str
    text: str
    page: int | None = None
    paragraph_index: int | None = None
    table_index: int | None = None
    row: int | None = None
    column: int | None = None
    style: dict[str, Any] = field(default_factory=dict)
    context_before: str = ""
    context_after: str = ""

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)


@dataclass
class RequirementSource:
    id: str
    path: str
    type: str
    blocks: list[SourceBlock]
    warnings: list[str] = field(default_factory=list)

    def metadata(self) -> dict[str, Any]:
        return {"id": self.id, "path": self.path, "type": self.type, "warnings": list(self.warnings)}


def _read_text(path: Path) -> tuple[str, list[str]]:
    warnings = []
    for encoding in ("utf-8", "gb18030", "latin-1"):
        try:
            return path.read_text(encoding=encoding), warnings
        except UnicodeDecodeError:
            warnings.append(f"{path.name}: 使用 {encoding} 解码失败")
    raise ValueError(f"无法读取文本文件: {path}")


def _with_context(blocks: list[SourceBlock]) -> list[SourceBlock]:
    for index, block in enumerate(blocks):
        block.context_before = blocks[index - 1].text if index > 0 else ""
        block.context_after = blocks[index + 1].text if index + 1 < len(blocks) else ""
    return blocks


def _extract_text_source(path: Path, source_id: str) -> RequirementSource:
    text, warnings = _read_text(path)
    blocks = []
    for index, line in enumerate(text.splitlines(), start=1):
        if not line.strip():
            continue
        block_type = "paragraph"
        style: dict[str, Any] = {"line_number": index}
        if path.suffix.lower() == ".md":
            stripped = line.lstrip()
            heading_level = len(stripped) - len(stripped.lstrip("#"))
            if heading_level and stripped[heading_level : heading_level + 1] == " ":
                block_type = "heading"
                style["heading_level"] = heading_level
        blocks.append(
            SourceBlock(
                id=f"{source_id}_line_{index:04d}",
                source_id=source_id,
                type=block_type,
                text=line.strip(),
                paragraph_index=len(blocks),
                style=style,
            )
        )
    return RequirementSource(source_id, str(path), path.suffix.lower().lstrip("."), _with_context(blocks), warnings)


def _run_style(run) -> dict[str, Any]:
    from docx.oxml.ns import qn

    style: dict[str, Any] = {}
    if run.font.name:
        style["font_name"] = run.font.name
    if run.font.size:
        style["font_size_pt"] = round(run.font.size.pt, 3)
    if run.bold is not None:
        style["bold"] = bool(run.bold)
    if run.italic is not None:
        style["italic"] = bool(run.italic)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is not None:
        east_asia = rfonts.get(qn("w:eastAsia"))
        ascii_font = rfonts.get(qn("w:ascii"))
        if east_asia:
            style["east_asia_font"] = east_asia
        if ascii_font:
            style["ascii_font"] = ascii_font
    return style


def _paragraph_style(paragraph) -> dict[str, Any]:
    style: dict[str, Any] = {"style_name": paragraph.style.name if paragraph.style else None}
    if paragraph.alignment is not None:
        style["alignment"] = str(paragraph.alignment)
    fmt = paragraph.paragraph_format
    for name in ("space_before", "space_after", "first_line_indent", "left_indent", "right_indent"):
        value = getattr(fmt, name)
        if value is not None:
            style[name] = {"pt": round(value.pt, 3), "cm": round(value.cm, 3)}
    if fmt.line_spacing is not None:
        value = fmt.line_spacing
        style["line_spacing"] = round(value, 3) if isinstance(value, float) else {"pt": round(value.pt, 3)}
    run_styles = []
    for run in paragraph.runs:
        if not run.text:
            continue
        run_style = _run_style(run)
        if run_style:
            run_styles.append(run_style)
    if run_styles:
        style["runs"] = run_styles
    return style


def _extract_docx(path: Path, source_id: str) -> RequirementSource:
    from docx import Document

    document = Document(str(path))
    blocks = []
    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        if text:
            blocks.append(
                SourceBlock(
                    id=f"{source_id}_paragraph_{index:04d}",
                    source_id=source_id,
                    type="paragraph",
                    text=text,
                    paragraph_index=index,
                    style=_paragraph_style(paragraph),
                )
            )
    for table_index, table in enumerate(document.tables):
        for row_index, row in enumerate(table.rows):
            for column_index, cell in enumerate(row.cells):
                text = cell.text.strip()
                if not text:
                    continue
                blocks.append(
                    SourceBlock(
                        id=f"{source_id}_table_{table_index:03d}_r{row_index:03d}_c{column_index:03d}",
                        source_id=source_id,
                        type="table_cell",
                        text=text,
                        table_index=table_index,
                        row=row_index,
                        column=column_index,
                    )
                )
    for section_index, section in enumerate(document.sections):
        section_style = {
            "orientation": str(section.orientation),
            "page_width_cm": round(section.page_width.cm, 3),
            "page_height_cm": round(section.page_height.cm, 3),
            "margin_top_cm": round(section.top_margin.cm, 3),
            "margin_bottom_cm": round(section.bottom_margin.cm, 3),
            "margin_left_cm": round(section.left_margin.cm, 3),
            "margin_right_cm": round(section.right_margin.cm, 3),
            "header_distance_cm": round(section.header_distance.cm, 3),
            "footer_distance_cm": round(section.footer_distance.cm, 3),
            "different_first_page": bool(section.different_first_page_header_footer),
        }
        blocks.append(
            SourceBlock(
                id=f"{source_id}_section_{section_index:03d}",
                source_id=source_id,
                type="document_settings",
                text=f"DOCX第{section_index + 1}节页面设置（模板实际样式，非文字要求）",
                style=section_style,
            )
        )
        for area_name, paragraphs in (
            ("header", section.header.paragraphs),
            ("footer", section.footer.paragraphs),
        ):
            for index, paragraph in enumerate(paragraphs):
                text = paragraph.text.strip()
                if text:
                    blocks.append(
                        SourceBlock(
                            id=f"{source_id}_{area_name}_{section_index:03d}_{index:03d}",
                            source_id=source_id,
                            type=area_name,
                            text=text,
                            paragraph_index=index,
                            style=_paragraph_style(paragraph),
                        )
                    )
    return RequirementSource(source_id, str(path), "docx", _with_context(blocks))


def _extract_pdf(path: Path, source_id: str) -> RequirementSource:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    warnings = []
    blocks = []
    for page_index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if not text.strip():
            warnings.append(f"{path.name}: 第 {page_index} 页未提取到文本，可能需要 OCR")
        for line_index, line in enumerate(text.splitlines(), start=1):
            if line.strip():
                blocks.append(
                    SourceBlock(
                        id=f"{source_id}_page_{page_index:04d}_line_{line_index:04d}",
                        source_id=source_id,
                        type="paragraph",
                        text=line.strip(),
                        page=page_index,
                        paragraph_index=line_index - 1,
                    )
                )
    if not blocks:
        warnings.append(f"{path.name}: PDF文本为空，当前版本未执行OCR")
    return RequirementSource(source_id, str(path), "pdf", _with_context(blocks), warnings)


def extract_requirement_source(path: Path, index: int = 1) -> RequirementSource:
    path = path.expanduser().resolve()
    if not path.is_file():
        raise ValueError(f"格式要求文件不存在或不是文件: {path}")
    if path.suffix.lower() not in SUPPORTED_SUFFIXES:
        raise ValueError(f"暂不支持的格式要求文件类型: {path.suffix}")
    source_id = f"source_{index:02d}"
    if path.suffix.lower() in {".txt", ".md", ".json"}:
        return _extract_text_source(path, source_id)
    if path.suffix.lower() == ".docx":
        return _extract_docx(path, source_id)
    return _extract_pdf(path, source_id)


def extract_requirement_sources(paths: list[Path]) -> list[RequirementSource]:
    if not paths:
        raise ValueError("至少需要一个格式要求文件")
    return [extract_requirement_source(path, index) for index, path in enumerate(paths, start=1)]


def write_source_model(path: Path, sources: list[RequirementSource]) -> None:
    data = {
        "sources": [source.metadata() for source in sources],
        "blocks": [block.to_dict() for source in sources for block in source.blocks],
    }
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
