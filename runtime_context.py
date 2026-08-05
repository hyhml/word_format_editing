"""Library helpers for AI-resolved, format-only runtime context.

This module deliberately has no CLI.  A skill uses it to build a minimal
evidence request from a target DOCX and to validate the AI's structured reply.
"""

from __future__ import annotations

import hashlib
import json
import re
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from jsonschema import Draft202012Validator

from format_ontology import TEMPLATE_FIELD_RUNTIME_INPUTS


RUNTIME_CONTEXT_SCHEMA_VERSION = "1.0.0"
SCHEMA_PATH = Path(__file__).resolve().parent / "schemas" / "runtime_context.schema.json"
MIN_CONFIDENCE = 0.8
MAX_EVIDENCE_BLOCKS = 30

FIELD_VALUE_CONTRACTS: dict[str, list[str]] = {
    "document.degree_type": ["doctoral", "master", "professional_master"],
    "document.project_type": ["thesis", "design"],
}

DEGREE_NAME_BY_TYPE = {
    "doctoral": "博士",
    "master": "硕士",
    "professional_master": "硕士",
}


def load_runtime_context_schema() -> dict[str, Any]:
    return json.loads(SCHEMA_PATH.read_text(encoding="utf-8"))


def _actions(rule: dict[str, Any]) -> Iterable[dict[str, Any]]:
    for action in rule.get("properties", {}).values():
        if not isinstance(action, dict):
            continue
        yield action
        if action.get("action") == "conditional":
            yield from (case for case in action.get("cases", []) if isinstance(case, dict))
            fallback = action.get("fallback")
            if isinstance(fallback, dict):
                yield fallback


def required_runtime_inputs(spec: dict[str, Any]) -> list[str]:
    """Return only runtime fields actually referenced by templates or conditions."""
    required = set()

    def scan_template(value: Any) -> None:
        if not isinstance(value, dict) or not isinstance(value.get("segments"), list):
            return
        for segment in value["segments"]:
            if not isinstance(segment, dict):
                continue
            if segment.get("kind") == "field":
                runtime_input = TEMPLATE_FIELD_RUNTIME_INPUTS.get(segment.get("field"))
                if runtime_input and runtime_input in FIELD_VALUE_CONTRACTS:
                    required.add(runtime_input)
            elif segment.get("kind") == "choice":
                for option in segment.get("options", []):
                    scan_template(option)

    def scan_condition(condition: Any) -> None:
        if not isinstance(condition, dict):
            return
        for logical in ("all", "any"):
            for child in condition.get(logical, []):
                scan_condition(child)
        if "not" in condition:
            scan_condition(condition["not"])
        field = condition.get("field")
        if field in FIELD_VALUE_CONTRACTS:
            required.add(field)

    rules = [spec.get("document", {}), *(spec.get("targets", {}).values() if isinstance(spec.get("targets"), dict) else [])]
    for rule in rules:
        if not isinstance(rule, dict):
            continue
        for action in _actions(rule):
            scan_template(action.get("value"))
            scan_condition(action.get("when"))
    return sorted(required)


def _document_fingerprint(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _paper_blocks(path: Path) -> list[dict[str, Any]]:
    document = Document(path)
    blocks = []
    for index, paragraph in enumerate(document.paragraphs):
        text = re.sub(r"\s+", " ", paragraph.text).strip()
        if text:
            blocks.append({"id": f"paper_paragraph_{index:04d}", "type": "paragraph", "paragraph_index": index, "text": text})
    for table_index, table in enumerate(document.tables):
        for row_index, row in enumerate(table.rows):
            for column_index, cell in enumerate(row.cells):
                text = re.sub(r"\s+", " ", cell.text).strip()
                if text:
                    blocks.append(
                        {
                            "id": f"paper_table_{table_index:03d}_r{row_index:03d}_c{column_index:03d}",
                            "type": "table_cell",
                            "table_index": table_index,
                            "row": row_index,
                            "column": column_index,
                            "text": text,
                        }
                    )
    return blocks


def _degree_evidence(blocks: list[dict[str, Any]]) -> list[dict[str, Any]]:
    patterns = (
        (r"专业学位|工程硕士|工商管理硕士|公共管理硕士|MBA|MPA", 16),
        (r"博士", 14),
        (r"硕士", 14),
        (r"学位(?:类别|类型|论文)", 8),
        (r"培养类别|专业类别|申请学位", 6),
    )
    scored = []
    for order, block in enumerate(blocks):
        score = sum(weight for pattern, weight in patterns if re.search(pattern, block["text"], re.IGNORECASE))
        if score and order < 30:
            score += 2
        if score:
            scored.append((score, -order, block))
    scored.sort(key=lambda item: (item[0], item[1]), reverse=True)
    selected = [item[2] for item in scored[:MAX_EVIDENCE_BLOCKS]]
    return selected


def _project_type_evidence(blocks: list[dict[str, Any]]) -> list[dict[str, Any]]:
    patterns = (
        (r"本科毕业设计", 18),
        (r"本科毕业论文", 18),
        (r"选题类型|项目类型|论文类型", 10),
        (r"毕业设计|毕业论文", 8),
    )
    scored = []
    for order, block in enumerate(blocks):
        score = sum(weight for pattern, weight in patterns if re.search(pattern, block["text"], re.IGNORECASE))
        if score and order < 30:
            score += 2
        if score:
            scored.append((score, -order, block))
    scored.sort(key=lambda item: (item[0], item[1]), reverse=True)
    return [item[2] for item in scored[:MAX_EVIDENCE_BLOCKS]]


def build_runtime_context_request(spec: dict[str, Any], paper_path: Path) -> dict[str, Any]:
    """Build the bounded evidence packet that the skill gives to the AI."""
    paper_path = paper_path.resolve()
    if paper_path.suffix.lower() != ".docx":
        raise ValueError("运行时格式上下文目前只读取待修改的 .docx 文件")
    required_fields = required_runtime_inputs(spec)
    blocks = _paper_blocks(paper_path) if required_fields else []
    evidence = []
    if "document.degree_type" in required_fields:
        evidence.extend(_degree_evidence(blocks))
    if "document.project_type" in required_fields:
        evidence.extend(_project_type_evidence(blocks))
    unique_evidence = {block["id"]: block for block in evidence}
    return {
        "schema_version": RUNTIME_CONTEXT_SCHEMA_VERSION,
        "document_fingerprint": _document_fingerprint(paper_path),
        "required_fields": required_fields,
        "allowed_values": {field: FIELD_VALUE_CONTRACTS[field] for field in required_fields},
        "unknown_policy": "证据不足、冲突或低置信度时使用status=unknown且value=null",
        "instruction": (
            "只判断格式规则所需字段；仅使用evidence_blocks；不得评价或改写论文内容。"
            "证据冲突、证据不足或无法确定时返回unknown。"
        ),
        "evidence_blocks": list(unique_evidence.values()),
    }


def validate_runtime_context_response(request: dict[str, Any], response: Any) -> dict[str, Any]:
    """Validate identity, field scope, evidence references, and confidence."""
    errors = []
    schema = load_runtime_context_schema()
    for issue in sorted(Draft202012Validator(schema).iter_errors(response), key=lambda item: list(item.absolute_path)):
        errors.append({"error_type": "schema_error", "path": ".".join(map(str, issue.absolute_path)), "message": issue.message})
    if not isinstance(response, dict):
        return {"status": "failed", "errors": errors}
    if response.get("document_fingerprint") != request.get("document_fingerprint"):
        errors.append({"error_type": "document_mismatch", "path": "document_fingerprint", "message": "AI响应不属于当前Word文件"})
    expected_fields = set(request.get("required_fields", []))
    received_fields = set(response.get("values", {})) if isinstance(response.get("values"), dict) else set()
    if received_fields != expected_fields:
        errors.append(
            {
                "error_type": "field_scope_mismatch",
                "path": "values",
                "message": "AI必须且只能返回format_spec实际需要的运行时字段",
                "expected": sorted(expected_fields),
                "received": sorted(received_fields),
            }
        )
    evidence_ids = {block["id"] for block in request.get("evidence_blocks", [])}
    for field, resolution in response.get("values", {}).items() if isinstance(response.get("values"), dict) else []:
        if not isinstance(resolution, dict):
            continue
        cited_ids = resolution.get("evidence_ids", [])
        if not isinstance(cited_ids, list) or not all(isinstance(item, str) for item in cited_ids):
            continue
        unknown_ids = sorted(set(cited_ids) - evidence_ids)
        if unknown_ids:
            errors.append(
                {
                    "error_type": "unknown_evidence",
                    "path": f"values.{field}.evidence_ids",
                    "message": "AI引用了请求中不存在的证据",
                    "received": unknown_ids,
                }
            )
    return {"status": "failed" if errors else "success", "errors": errors}


def normalize_runtime_context(
    request: dict[str, Any], response: Any, minimum_confidence: float = MIN_CONFIDENCE
) -> dict[str, Any]:
    """Return safe context; low confidence becomes unknown instead of a guess."""
    validation = validate_runtime_context_response(request, response)
    if validation["errors"]:
        return {"status": "blocked", "values": {}, "errors": validation["errors"]}
    values = {}
    warnings = []
    for field, resolution in response["values"].items():
        item = dict(resolution)
        if item["status"] == "resolved" and item["confidence"] < minimum_confidence:
            warnings.append({"field": field, "warning_type": "low_confidence", "received": item["confidence"]})
            item = {
                "status": "unknown",
                "value": None,
                "confidence": item["confidence"],
                "evidence_ids": item["evidence_ids"],
                "reason": "置信度低于阈值，保持原格式",
            }
        values[field] = item
    return {
        "status": "warn" if warnings or any(item["status"] == "unknown" for item in values.values()) else "success",
        "document_fingerprint": request["document_fingerprint"],
        "values": values,
        "warnings": warnings,
        "errors": [],
    }


def template_context(runtime_context: dict[str, Any], institution_name: str | None = None) -> dict[str, str]:
    """Translate validated runtime values into concrete template fields."""
    result = {}
    degree = runtime_context.get("values", {}).get("document.degree_type", {})
    if degree.get("status") == "resolved" and degree.get("value") in DEGREE_NAME_BY_TYPE:
        result["degree_name"] = DEGREE_NAME_BY_TYPE[degree["value"]]
    if institution_name:
        result["institution_name"] = institution_name
    return result
